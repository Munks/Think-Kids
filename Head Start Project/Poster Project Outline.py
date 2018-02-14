from docx import Document
from docx.shared import Pt

path = r'\\Cifs2\thinkkid$\Research\Chris\Head Start Project\Poster'
doc_name = '\Head Start Poster Outline.docx'


def para(text):
    return document.add_paragraph(text, style='Normal')


def bullet(text):
    document.add_paragraph(text, style='List Bullet')


"""
Creating a the Poster Project Outline
"""
document = Document()
times = 'Times New Roman'

# Setting up the fonts
font = document.styles['Normal'].font
font.name = times
font.size = Pt(12)

font = document.styles['Title'].font
font.name = times
font.size = Pt(16)


p = para('Questions to answer for the Poster Presentation')
bullet('What data do we have to answer this question?')
bullet('How are we going to present this data?')

p = para('\nDemographics of the group')
bullet('We have age and race of the parent, age of the child, and education level of the parent. We also have the ACES scores of the parents.')
bullet('We can present the ages of the parents and children as bar graphs and the race and education levels as pie charts. The ACES can be presented as an ordered bar graph?')

p = para('\nFeasibility of running parent group for preschoolers')
bullet('We have the drop out rate of parents between time 1, 2 and 3.')
bullet('We should present this data as a bar graph')

p = para('\nOutcomes for those who attended')
bullet('We have scored the CPS-AIM and PCRI across all time points. We can score the DECA, DERS, PRFCS and the PSDQ, but I am unsure the degree of relivent data that we will recieve from them. We are also missing data on 4 out of the 6 participants with the DECA for time 3')
bullet('This data is most likely best presented as bar graphs as well.')

document.save(path + doc_name)
