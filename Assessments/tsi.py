from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def font_style(name, size, font_type):
    font_type = font_type
    font = document.styles[name].font
    font.name = font_type
    font.size = Pt(size)


def set_col_widths(table, first_col, second_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(second_col),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def table_font_size(size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(11)


document = Document('TK Template.docx')

styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

font_style('Normal', 12, 'Times New Roman')
font_style('Body Text', 14, 'Times New Roman')
font_style('Body Text 2', 18, 'Times New Roman')
font_style('Title', 16, 'Times New Roman')

questions = [
    'It’s hard for me to stay focused on things that I need to\n',
    'It’s hard for me to remember the steps or directions I need to get \
things done',
    'It’s hard for me to keep track of time to get places and do things on \
time',
    'I have a hard time understanding what other people are trying to tell me',
    'I have a hard time telling people how I feel\n',
    'I have a hard time telling people what I am thinking\n',
    'It’s hard for me to settle down when I am hyped up\n',
    'It’s hard for me to get my exergy level up when I need to\n',
    'It’s hard to control my worries\n',
    'I have a hard time thinking straight when I am feeling frustrated',
    'I have a hard time handling things when I am feeling disappointed',
    'It’s hard for me to stop and think before I say or do things',
    'I don’t do well in new or unexpected situations\n',
    'I have a hard time when my plans or schedule changes\n',
    'It’s hard for me to stop one thing I’m doing and start up a new thing',
    'It’s hard for me to think of more than one way to solve a problem',
    'I tend to take things too personally\n',
    'I tend to exaggerate, make too big of a deal about things or think that \
things are worse than they are',
    'I don’t usually notice or understand people’s facial expression or tone \
of voice when they are talking to me',
    'It’s hard for me to tell what people think about me\n',
    'I’m not very good at talking to new people\n',
    'I have a hard time understanding how other people are feeling'
]

for x, i in enumerate(questions):
    if x == 11:
        document.add_page_break()
    if x == 0 or x == 11:
        p = document.add_paragraph('', style='Body Text 2')
        p.add_run('Youth Thinking Skills Inventory')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(
            rows=1, cols=5, style='Medium List 2 Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[1].paragraphs[0].add_run('Questions').bold = True
        hdr_cells[2].paragraphs[0].add_run('Never/ Rarely')
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].paragraphs[0].add_run('Sometimes')
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].paragraphs[0].add_run('Often/ Always')
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_col_widths(table, .1, 4, .5)
    table = document.add_table(rows=1, cols=5, style='Medium List 2 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(str(x + 1) + ')')
    hdr_cells[1].paragraphs[0].add_run(i)
    hdr_cells[2].paragraphs[0].add_run('0')
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[3].paragraphs[0].add_run('1')
    hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[4].paragraphs[0].add_run('2')
    hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(table, .1, 4, .5)


document.add_page_break()

p = document.add_paragraph('Scoring', 'Body Text 2')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Probably add description of scoring here

skills = {'Attention and Working Memory': '1, 2, 3',
          'Language and Communication': '4, 5, 6',
          'Emotion and Self-Regulation': '7, 8, 9, 10, 11, 12',
          'Cognitive Flexibility': '13, 14, 15, 16, 17, 18',
          'Social Thinking': '19, 20, 21, 22'}

p = document.add_paragraph('', 'Normal')
p.add_run('The TSI is scored across multiple skill areas. Each domain score \
is an average of all possible points within the domain. Add the items in each \
skill and place your answers in the table below.')

table = document.add_table(rows=1, cols=3, style='Medium List 2 Accent 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Domain').bold = True
hdr_cells[1].paragraphs[0].add_run('Question Numbers')
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('Total Score')
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table_font_size(11)
for key, value in skills.items():
    table = document.add_table(rows=1, cols=3, style='Medium List 2 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(key)
    hdr_cells[1].paragraphs[0].add_run(value)
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[2].paragraphs[0].add_run('__________')
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_font_size(11)


document.add_paragraph('\n\nThe following chart is used to see how the domains \
compare. Mark an X under each skill name to indicate the score.', 'Normal')

table = document.add_table(rows=1, cols=6, style='Medium List 2 Accent 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Average Score').bold = True
hdr_cells[1].paragraphs[0].add_run('Attention and Working Memory')
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('Language and Communication')
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].paragraphs[0].add_run('Emotion and Self-Regulation')
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[4].paragraphs[0].add_run('Cognitive Flexibility')
hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[5].paragraphs[0].add_run('Social Thinking')
hdr_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table_font_size(11)

for i in range(13):
    table = document.add_table(rows=1, cols=6, style='Medium List 2 Accent 1')
    hdr_cells = table.rows[0].cells
    if i == 0:
        hdr_cells[0].paragraphs[0].add_run(str(2) + ' - Big Problem')
    if i == 12:
        hdr_cells[0].paragraphs[0].add_run(str(0) + ' - No Problem')
    for i in range(5):
        hdr_cells[i + 1].paragraphs[0].add_run('-')
        hdr_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_font_size(11)

set_col_widths(table, 1.2, 1)

document.save('tsi.docx')
