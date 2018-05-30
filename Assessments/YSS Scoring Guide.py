from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches


def font_style(name, size, font_type):
    font_type = font_type
    font = document.styles[name].font
    font.name = font_type
    font.size = Pt(size)


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


document = Document('TK Template.docx')

styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

font_style('Normal', 12, 'Times New Roman')
font_style('Body Text', 10, 'Times New Roman')
font_style('Title', 16, 'Times New Roman')

questions = [
    ''
    'Overall, I am satisfied with the services I received.',
    'I helped to choose my services.',
    'I helped to choose my treatment goals.',
    'The people helping me stuck with me no matter what.',
    'I felt I had someone to talk to when I was in trouble.',
    'I participated in my own treatment.',
    'I received services that were right for me.',
    'The location of services was convenient for me.',
    'Services were available at times that were convenient for me.',
    'I got the help I wanted.',
    'I got as much help as I needed.',
    'Staff treated me with respect.',
    'Staff respected my religious / spiritual beliefs.',
    'Staff spoke with me in a way that I understood.',
    'Staff were sensitive to my cultural / ethnic background.',
    'I am better at handling daily life.',
    'I get along better with family members.',
    'I get along better with friends and other people.',
    'I am doing better in school and / or work.',
    'I am better able to cope when things go wrong.',
    'I am satisfied with my family life right now.',
    'I am better able to do things I want to do.',
    'I know people who will listen and understand me when I need to talk.',
    'I have people that I am comfortable talking with about my problem(s).',
    'In a crisis, I would have the support I need from family or friends.',
    'I have people with whom I can do enjoyable things.',
    'What has been the most helpful thing about the services you received \
over the last 6 months?',
    'What would improve the services here?'
]

document.add_heading('Youth Services Survey', 0)

p = document.add_paragraph('', style='Normal')
p.add_run(
    '\tThe YSS includes a total of 26 items, and two additional comment items \
shown in Table 1. The 26 items are divided into seven domains, show in \
Table 2.'
)

p = document.add_paragraph('', 'No Spacing')
p.add_run('Table 1\n').bold = True

for x, y in enumerate(questions):
    table = document.add_table(rows=1, cols=2, style='Medium List 2 Accent 1')
    if x == 0:
        hdr_cell = table.rows[0].cells
        hdr_cell[1].paragraphs[0].add_run('The questions are as follows:')
    else:
        row_cells = table.rows[0].cells
        row_cells[0].paragraphs[0].add_run(str(x + 1) + ')')
        row_cells[1].paragraphs[0].add_run(y)
    set_col_widths(table, .1, 6.9)

document.add_page_break()

domains_dict = {
    'Domain': 'Survey Item Numbers',
    'General Satisfaction': '1, 4, 5, 7, 10, 11',
    'Participation in Treatment Planning': '2, 3, 6',
    'Access': '8, 9',
    'Cultural Sensitivity': '12, 13, 14, 15',
    'Social Connectedness': '23, 24, 25, 26',
    'Outcomes': '16, 17, 18, 19, 20, 21',
    'Functioning': '16, 17, 18, 19, 20, 22'
}

p = document.add_paragraph('', 'No Spacing')
p.add_run('Table 2').bold = True
t = document.add_table(rows=1, cols=2, style='Medium List 1 Accent 5')
for key, item in domains_dict.items():
    row_cells = t.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = item
set_col_widths(t, 3.25, 3.25)
document.add_paragraph(
    '\n\tThe content of the domains in the YSS instrument has been designed \
for the child mental health population. Each item on the YSS is answered \
using a Likert scale ranging from one (strongly disagree) to five \
(strongly agree). Items in a domain are summed and divided by the total \
number of items, and scores greater than 3.5 are reported in the positive \
range for the domain. Cases with domains where more than one-third of \
items are missing are not included in the final analysis. Additionally, \
the survey includes two questions that ask cosumers to share 1) what has \
been more helpful about the services and 2) what would improve services.'
)

document.save('YSS Scoring Guide.docx')
