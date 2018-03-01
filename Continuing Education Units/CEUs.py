import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import glob
import win32com.client


"""
Program for creating CEU PDF's for Think:Kids Teir 1 and Teir 2 trainings
"""

path = os.path.expanduser('~\Desktop')
# Functions


def intro(document):
    p = document.add_paragraph('\n\n\n', 'Normal')
    p.alignment = centered
    p = document.add_paragraph('Continuing Education Certificate',
                               'Cert Title')
    p.alignment = centered
    document.add_paragraph('', 'Normal')
    p = document.add_paragraph('Presented to:', 'Presented to')
    p.alignment = centered
    p = document.add_paragraph(df.loc[i, 'First Name'].title() +
                               ' ' + df.loc[i, 'Last Name'].title(), 'Name')
    p.alignment = centered


def tier(document):
    p = document.add_paragraph('', 'cps')
    p.alignment = centered
    p.add_run(tier_num)
#    p.add_run('Tier ' + tier_question + ' Intensive Training')
    p = document.add_paragraph('', 'cps')
    p.alignment = centered
    if trainer_count == 1:
        trainers = 'Think:Kids Certified Trainer: ' + facilitators[0]
        p.add_run(trainers)
    if trainer_count == 2:
        trainers = 'Think:Kids Certified Trainers: ' + \
            facilitators[0] + ' and ' + facilitators[1]
        p.add_run(trainers)
    if trainer_count == 3:
        trainers = 'Think:Kids Certified Trainers: ' + \
            facilitators[0] + ', ' + facilitators[1] + \
            ' and ' + facilitators[2]
        p.add_run(trainers)
    p = document.add_paragraph('', 'Date1')
    p.alignment = centered
    p.add_run(date_of_training_start + ' - ' + date_of_training_end)
    p = document.add_paragraph('', 'cps')
    p.alignment = centered
    p.add_run(location)


def director_sig(document):
    p = document.add_paragraph('', 'Normal')
    p.add_run('Director, Think:Kids')
    p.add_run('\nDept. of Psychiatry')
    p.add_run('\nMassachusetts General Hospital')
    p.add_run('\n151 Merrimac Street, 5th Floor')
    p.add_run('\nBoston, MA. 02114')
    p.add_run('\nwww.thinkkids.org')


def stuart_sig_general(document):
    # For General, Social Work, Educator
    p = document.add_paragraph('\n\n', 'Normal')
    p.add_run('\t\t\t\t\t\t\t\t   ')
    p.add_run('      16      ').underline = True
    p.add_run('\t\t\t\t\t')
    p.add_run(date_of_issue).underline = True
    p.add_run('\nJ. Stuart Ablon, PhD')
    p.add_run('\t\t\t\t\tContinuing Education Hours')
    p.add_run('\t\t\t   Date Issued')
    director_sig(document)


def stuart_sig_lmhc(document):
    # For General, Social Work, Educator
    p = document.add_paragraph('\n\n', 'Normal')
    p.add_run('\t\t\t\t\t\t   ')
    p.add_run('      16      ').underline = True
    p.add_run('\t\t\t\t')
    p.add_run('   1   ').underline = True
    p.add_run('\t\t\t')
    p.add_run(date_of_issue).underline = True
    p.add_run('\nJ. Stuart Ablon, PhD')
    p.add_run('\t\t\tContinuing Education Hours')
    p.add_run('\t        Category')
    p.add_run('\t\t\t   Date Issued')
    director_sig(document)


def sig_psych(document):
    # For Psychiatry CEUs
    p = document.add_paragraph('\n\n\n', 'Normal')
    p.add_run('\t\t\t\t\t\t   ')
    p.add_run('   16 Hours   ').underline = True
    p.add_run('\t\t')
    p.add_run(date_of_issue).underline = True
    p.add_run('\t\t\t')
    p.add_run('\nJ. Stuart Ablon, PhD')
    p.add_run('\t\t\tContinuing Education Hours')
    p.add_run('\t   Date Issued')
    p.add_run('\t\t\t         Dr. Susan Sprich, Ph.D.')
    p.add_run('\nDirector, Think:Kids')
    p.add_run(
        '\t\t\t\t\t\t\t\t\t     Director for Postgraduate Psychology Education')
    p.add_run('\nDept. of Psychiatry')
    p.add_run('\t\t\t\t\t\t\t\t\t\t\t       Postgraduate Medical Education')
    p.add_run('\nMassachusetts General Hospital')
    p.add_run(
        '\t\t\t\t\t\t          Massachusetts General Hospital, Dept. of Psychiatry')
    p.add_run('\n151 Merrimac Street, 5th Floor')
    p.add_run('\t\t\t\t\t\t\t\t\t         One Bowdoin Square, 7th floor')
    p.add_run('\nBoston, MA. 02114')
    p.add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t      Boston, MA 02114')
    p.add_run('\nwww.thinkkids.org')


def license_num(document):
    p = document.add_paragraph('License #_______________', 'Normal')
    p.alignment = centered


def mental_health_cert(document):
    p = document.add_paragraph(
        '\nThis training has been approved for 16 continuing education hours' +
        ' by MMCEP and is approved as a continuing education activity for' +
        ' licensed mental health clinicians. Authorization Number: 17-0796\n'
    )
    p.alignment = centered
    stuart_sig_lmhc(document)


def social_work_cert(document):
    p = document.add_paragraph(
        '\nThis training has been approved for 16 Approved Entity Continuing' +
        ' Education hours for relicensure,\nin accordance with 258 CMR.' +
        ' Collaborative of NASW and the Boston College and Simmons School ' +
        'of Social Work.\nAuthorization Number: D 72498'
    )
    p.alignment = centered
    stuart_sig_general(document)


def educators_cert(document):
    p = document.add_paragraph(
        '\nThink:Kids is a registered Professional Development Point(PDP) ' +
        'Provider with the Massachusetts Department of Elementary and ' +
        'Secondary Education.\nThis training has been approved for 16 PDPs.\n'
    )
    p.alignment = centered
    stuart_sig_general(document)


def general_cert(document):
    document.add_paragraph('\n\n\n\n')
    stuart_sig_general(document)


def psych_cert(document):
    p = document.add_paragraph(
        '\nThis training has been approved for 16 continuing education ' +
        'hours by the Massachusetts General Hospital, Department of ' +
        'Psychiatry Continuing Education Division\n')
    p.alignment = centered
    sig_psych(document)


def build_start(document):
    intro(document)
    if df.loc[i, discipline] == lmhc or df.loc[i, discipline] == social_work:
        license_num(document)
    tier(document)


# User Interface


print("\n\n\t\t* * * * * * *   *    *      *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *  *")
print("\t\t      *              **")
print("\t\t      *         *    *  *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *      *")
print("\n\t\t     CEU Certificate Magic!\n\n")

trainer = "Think:Kids Certified Trainer(s)"

setup = True

while setup:
    tier_num = input(
        "Please Enter the Type of Tier Training? (Please enter 1 or 2)\n\t")
    date_of_training_start = input(
        "\nWhat Date did the training begin on? " +
        "(Please type a full date e.g. January 1)\n\t").capitalize()
    date_of_training_end = input(
        "\nWhat Date did the training end on? " +
        "(Please type a full date e.g. January 1st, 2018)\n\t").capitalize()
    trainer_count = int(input("\nHow many " + trainer +
                              " were there? (enter a number 1-3)\n\t"))

    facilitators = []
    if trainer_count > 0:
        first_facilitator = input("\nWho was the First " + trainer + "?\n\t")
        facilitators.append(first_facilitator)
    if trainer_count > 1:
        second_facilitator = input("\nWho was the Second " + trainer + "?\n\t")
        facilitators.append(second_facilitator)
    if trainer_count > 2:
        third_facilitator = input("\nWho was the Third " + trainer + "?\n\t")
        facilitators.append(third_facilitator)

    location = input("\nPlease enter the location of the training " +
                     "e.g Boston, MA\n\t")
    date_of_issue = input("\nPlease enter the Date of Issue that you would " +
                          "like to appear on the Certificate " +
                          "e.g. January 1st, 2018\n\t").capitalize()
    print("\n\n\n\nYou've entered:")
    print("\tThis was a: \t\t\tTier " + tier_num + " Training")
    print("\tThe Training occured: \t\t" + date_of_training_start +
          " - " + date_of_training_end)
    print("\tYour Facilitators were:")
    for i in facilitators:
        print("\t\t\t\t\t" + i)
    print("\tThis training happened in:\t" + location)
    print("\tThe date of issue is:\t\t" + date_of_issue + '\n\n')
    print("\t***\tPlease double check this information Very Carefully!\t***")
    print("\nIt will come out exactly as it appears here in the CEU PDFs")
    print("\nCheck for spelling mistakes and missing comas in the dates.")
    print("\nCheck to make sure the appropriate things are capitalized.")
    print("\nMake sure the facilitators are in the correct order.")
    setup_correct = input("\n\nIs this correct? Type Yes to continue" +
                          " or No to re-enter this information\n\t").title()
    if setup_correct == 'Yes':
        setup = False
    elif setup_correct == 'No':
        setup = True

# Make folder system on desktop
ceus_folder = path + "\CEUs for Tier " + \
    tier_num + " " + location
data_folder = ceus_folder + "\Original Spreadsheet"
social_work_folder = ceus_folder + r'\Social Worker'
lmhc_folder = ceus_folder + r'\Mental Health Clinician'
educators_folder = ceus_folder + r'\Mass Educator'
general_folder = ceus_folder + r'\General Attendance Cert'
psychology_folder = ceus_folder + r'\Psychologist'

if not os.path.exists(ceus_folder):
    os.makedirs(ceus_folder)
if not os.path.exists(data_folder):
    os.makedirs(data_folder)

print("\n\n\n\n\n\t***\tLets pause for a second\t***")
print("\n\nThere is a new folder on your Desktop called:")
print("\tCEUs for Tier " + tier_num + " " + location)
print("\nGo and open it." +
      " Everything that this program makes will go in this folder.")
print("\nThe only thing there right now is a folder called" +
      " Original Spreadsheet.")
print("\nDownload the Excel file from Eventbrite and place it into the" +
      " Original Spreadsheet folder.")
print("\nMake sure that when you Download the Excel file from" +
      " Eventbrite it has the following column names:")
print("\t'First Name'")
print("\t'Last Name'")
print("\t'Please indicate discipline(s) for which you wish to " +
      "receive CEUs/Certificate of Attendance'")
print("\nCheck to make sure that the 'First Name' and 'Last Name' are " +
      "properly capitalized in the excel")
input("\n\nWhen you are have finished press Enter!")


df = pd.DataFrame({'A': []})
while df.empty:
    folder = glob.glob(data_folder + '\*.xlsx')
    for file in folder:
        df = pd.read_excel(file)
    if df.empty:
        input("\nPlease go place the Excel from Eventbrite into the Original Spreadsheet folder. Press enter when finished")


# Change to -16 when .exe
template_general = os.path.abspath(
    __file__)[:-7] + r'Templates\template_general.docx'
template_psych = os.path.abspath(
    __file__)[:-7] + r'Templates\template_psych.docx'


tier_num = 'Tier ' + tier_num
tier_num = tier_num + ' Intensive Training in\nCollaborative Problem Solving®'


discipline = 'Please indicate discipline(s) for which you wish to receive CEUs/Certificate of Attendance'
social_work = 'I am a Licensed Social Worker'
educator = 'I am a Massachusetts Educator'
general = 'Other, I will need a general certificate of attendance.'
parent = 'I am a parent; I will not need CEUs'
lmhc = 'I am a Licensed Mental Health Counselor'
psych = 'I am a Psychologist'


# Run Program

word = win32com.client.DispatchEx("Word.Application")
word.Visible = False


for i in df.index:
    if df.loc[i, discipline] == lmhc:
        if not os.path.exists(lmhc_folder + '\documents'):
            os.makedirs(lmhc_folder + '\documents')
        document = Document(template_general)
        centered = WD_PARAGRAPH_ALIGNMENT.CENTER
        left = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(df.loc[i, 'First Name'].title() +
              ' ' + df.loc[i, 'Last Name'].title())
        print(df.loc[i, discipline])
        build_start(document)
        mental_health_cert(document)
        doc_folder = lmhc_folder + '\documents'
        doc_name = doc_folder + r'\\' + df.loc[i, 'Last Name'] + '.docx'
        pdf_name = lmhc_folder + r'\\' + df.loc[i, 'Last Name'] + '.pdf'
        document.save(doc_name)
        doc = word.Documents.Open(doc_name)
        doc.SaveAs(pdf_name, FileFormat=17)
        doc.Close()

    if df.loc[i, discipline] == social_work:
        if not os.path.exists(social_work_folder + '\documents'):
            os.makedirs(social_work_folder + '\documents')
        document = Document(template_general)
        centered = WD_PARAGRAPH_ALIGNMENT.CENTER
        left = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(df.loc[i, 'First Name'].title() +
              ' ' + df.loc[i, 'Last Name'].title())
        print(df.loc[i, discipline])
        build_start(document)
        social_work_cert(document)
        doc_folder = social_work_folder + '\documents'
        doc_name = doc_folder + r'\\' + df.loc[i, 'Last Name'] + '.docx'
        pdf_name = social_work_folder + r'\\' + df.loc[i, 'Last Name'] + '.pdf'
        document.save(doc_name)
        doc = word.Documents.Open(doc_name)
        doc.SaveAs(pdf_name, FileFormat=17)
        doc.Close()

    if df.loc[i, discipline] == educator:
        if not os.path.exists(educators_folder + '\documents'):
            os.makedirs(educators_folder + '\documents')
        document = Document(template_general)
        centered = WD_PARAGRAPH_ALIGNMENT.CENTER
        left = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(df.loc[i, 'First Name'].title() +
              ' ' + df.loc[i, 'Last Name'].title())
        print(df.loc[i, discipline])
        build_start(document)
        educators_cert(document)
        doc_folder = educators_folder + '\documents'
        doc_name = doc_folder + r'\\' + df.loc[i, 'Last Name'] + '.docx'
        pdf_name = educators_folder + r'\\' + df.loc[i, 'Last Name'] + '.pdf'
        document.save(doc_name)
        doc = word.Documents.Open(doc_name)
        doc.SaveAs(pdf_name, FileFormat=17)
        doc.Close()

    if df.loc[i, discipline] == general or df.loc[i, discipline] == parent:
        if not os.path.exists(general_folder + '\documents'):
            os.makedirs(general_folder + '\documents')
        document = Document(template_general)
        centered = WD_PARAGRAPH_ALIGNMENT.CENTER
        left = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(df.loc[i, 'First Name'].title() +
              ' ' + df.loc[i, 'Last Name'].title())
        print(df.loc[i, discipline])
        build_start(document)
        general_cert(document)
        doc_folder = general_folder + '\documents'
        doc_name = doc_folder + r'\\' + df.loc[i, 'Last Name'] + '.docx'
        pdf_name = general_folder + r'\\' + df.loc[i, 'Last Name'] + '.pdf'
        document.save(doc_name)
        doc = word.Documents.Open(doc_name)
        doc.SaveAs(pdf_name, FileFormat=17)
        doc.Close()

    if df.loc[i, discipline] == psych:
        if not os.path.exists(psychology_folder + '\documents'):
            os.makedirs(psychology_folder + '\documents')
        document = Document(template_psych)
        centered = WD_PARAGRAPH_ALIGNMENT.CENTER
        left = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(df.loc[i, 'First Name'].title() +
              ' ' + df.loc[i, 'Last Name'].title())
        print(df.loc[i, discipline])
        build_start(document)
        psych_cert(document)
        doc_folder = psychology_folder + '\documents'
        doc_name = doc_folder + r'\\' + df.loc[i, 'Last Name'] + '.docx'
        pdf_name = psychology_folder + r'\\' + df.loc[i, 'Last Name'] + '.pdf'
        document.save(doc_name)
        doc = word.Documents.Open(doc_name)
        doc.SaveAs(pdf_name, FileFormat=17)
        doc.Close()
        # Add PDF

word.Quit()
print("\n\n\t\t***\tYou've Made CEU's Magic!!\t***")
