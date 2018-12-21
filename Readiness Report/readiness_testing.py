from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pandas as pd

source_text = pd.read_excel(
    'readiness_source_text.xlsx',
    sheet_name=None)

total_readiness = source_text['Total Readiness Table']
cps_aim_subscales_text = source_text['CPS-AIM Subscales']
readiness_survey = source_text['CPS Readiness Survey']

for i in readiness_survey.index:
    print(i)
    print(readiness_survey.iloc[i]['Paragraph'])
    print(readiness_survey.iloc[i]['Style'])
    print(readiness_survey.iloc[i]['Text'])


document = Document()


def add_document_text_from_excel(df):
    for i in df.index:
        if df.iloc[i]['Paragraph']:
            if df.iloc[i]['Paragraph'] == 'Skip':
                continue
            else:
                paragraph_style = str(df.iloc[i]['Paragraph'])
        if df.iloc[i]['Text']:
            text = str(df.iloc[i]['Text'])
        if df.iloc[i]['Style']:
            text_style = str(df.iloc[i]['Style'])
        if 'nan' not in paragraph_style:
            paragraph = document.add_paragraph('', style=paragraph_style)
        if text_style == 'bold':
            paragraph.add_run(text).bold = True
        elif text_style == 'italic':
            paragraph.add_run(text).italic = True
        elif text_style == 'underline':
            paragraph.add_run(text).underline = True
        else:
            paragraph.add_run(text)


add_document_text_from_excel(readiness_survey)

document.save('Test.docx')
