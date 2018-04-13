from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import datetime as dt
from redcap import Project
import sys
sys.path.append(r'C:\Python Programs\Think-Kids_Private')
import tokens as tk
import scoring as score


"""
Code to create a Readiness Report. The data is pulled from REDcap. Data is
retrieved any of the following three locations:

    "CPS Readiness Assessment (3-Part)"

    "CPS-AIM, Educators' Version"
        "CPS Readiness Assessment for Schools (1-Part)"

    "CPS-AIM Systems' Version"
        "CPS Readiness Assessment for Agencies (1-Part)"


This code is only currently working on the second option.
"""


# Functions

def meta_cell_content(df, row_name, column_name):
    """
    Retrieve the content of a given cell in the Data Dictionary downlaoded from
    REDcap
    """
    dict_name = df.loc[[row_name], [column_name]]
    dict_name = dict_name[column_name]
    item = []
    for i in dict_name:
        item.append(i)
    return item


def meta_dict(df, row_name):
    """
    Create a python dictionary from the Data Dictionary downloaded from REDcap
    df = Data Dictionary DataFrame Pandas Object
    row_name = Name of Row
    The column for this Dictionary is always 'select_choices_or_calculations'
    """
    column_name = 'select_choices_or_calculations'
    item = meta_cell_content(df, row_name, column_name)
    dic_key = []
    dic_val = []
    for i in item:
        item = i.split(' | ')
        for i in item:
            item = i.split(', ', 1)
            a = 0
            for i in item:
                if a == 0:
                    dic_key.append(int(i))
                    a += 1
                else:
                    dic_val.append(i)
    dict_name = dict(zip(dic_key, dic_val))
    return dict_name


def df_aim_columns(df, label, columns_list):
    for column in df.columns:
        if label in column[:-4]:
            name = column[:]
            columns_list.append(name)


def font_style(name, size, font_type):
    font_type = font_type
    font = document.styles[name].font
    font.name = font_type
    font.size = Pt(size)


def make_table(table_title, primary_dic, secondary_dic, table_count):
    table = document.add_table(rows=1, cols=4, style='Normal Table')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(table_title).bold = True
    headers = ['Freq', 'Percent', 'Cum.']
    for count, header in enumerate(headers):
        hdr_cells[count + 1].paragraphs[0].add_run(header).underline = True
    total = 0
    for key, value in primary_dic.items():
        row_cells = table.add_row().cells
        row_cells[0].text = secondary_dic[key]
        row_cells[1].text = str(value)
        total += int(value)
        row_cells[2].text = str(
            round(float((value / table_count) * 100), 0)) + '%'
        row_cells[3].text = str(
            round(float((total / table_count) * 100), 0)) + '%'
    set_col_widths(table, 3, 1)


# Defining the REDcap API Tokens
# Call these functions immediately before pulling the REDcap Data corresponding
# to these projects. Only the most recent call will be available.
# "CPS-AIM, Educators' Version"
tk.cps_aim_educators()
project = Project(tk.api_url, tk.api_token)
data_cps_aim_educators = project.export_records(format='df')
metadata_cps_aim_educators = project.export_metadata(format='df')

# User Interface Introduction
print("\n\n\t\t* * * * * * *   *    *      *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *  *")
print("\t\t      *              **")
print("\t\t      *         *    *  *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *      *")
print("\n\t\t\tReadiness Reports!!\n\n")

# Global Variables

current_date = str(dt.date.today().strftime('%m-%d-%Y'))


organization_dictionary_aim = meta_dict(metadata_cps_aim_educators, 'org')

# User Interface
print("The organizations listed in the CPS-AIM, Educators' Version are as \
follows.")
for key, value in organization_dictionary_aim.items():
    print(str(key) + '\t' + value)
organization_number_aim = int(input("Please enter the Organizations Number to \
generate a report on.\n\t"))
print('Thanks!\n\n\n\n')

# Retrieving and sorting the AIM Data from REDcap
data_cps_aim_educators.reset_index(inplace=True)
organization_data_aim = data_cps_aim_educators[
    data_cps_aim_educators['org'] == organization_number_aim
]
# Remove participants who did not answer at least 20 questions
organization_data_aim = organization_data_aim.dropna(thresh=20)

# Replace any N/A answers (9) with np.nan so that the data is not skewed
organization_data_aim = organization_data_aim.replace(9, np.nan)

# Variables from AIM Data Set

# Role variables
role_dictionary_aim = meta_dict(metadata_cps_aim_educators, 'role')
role_count_aim = organization_data_aim['role'].value_counts().to_dict()
role_survey_count_aim = organization_data_aim['role'].count()

# Get a count of how many participants responded to the survey and answered
# at least 20 questions.
survey_count_aim = organization_data_aim.record_id.count()


# Training Dictionary
training_columns = ['received_training', 'intro_training', 'tier_one_training']
training = organization_data_aim[training_columns]
training_dictionary = {
    'No Training': 0,
    'Not Sure': 0,
    'Introductory Training': 0,
    'Tier One Training': 0
}
training_total = 0
for i in training['received_training']:
    if i == 0:
        training_dictionary['No Training'] += 1
        training_total += 1
    if i == 2:
        training_dictionary['Not Sure'] += 1
        training_total += 1
for i in training['intro_training']:
    if i == 1:
        training_dictionary['Introductory Training'] += 1
        training_total += 1
for i in training['tier_one_training']:
    if i == 1:
        training_dictionary['Tier One Training'] += 1
        training_total += 1

# Sub School
sub_school_column = []
for x, i in enumerate(metadata_cps_aim_educators['branching_logic']):
    if '[org] = ' + "'" + str(organization_number_aim) + "'" in str(i):
        sub_school_column.append(metadata_cps_aim_educators.index[x])
sub_school_dict = {}
for i in sub_school_column:
    sub_school_dict = meta_dict(metadata_cps_aim_educators, i)
    sub_school_count = organization_data_aim[i].value_counts().to_dict()
    sub_school_survey_count = organization_data_aim[i].count()

# Creating the Report in a Word Docutment
document = Document('Readiness Report Template.docx')

styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

font_style('Normal', 12, 'Times New Roman')
font_style('Body Text', 14, 'Times New Roman')
font_style('Body Text 2', 18, 'Times New Roman')
font_style('Title', 16, 'Times New Roman')


# Title Page
paragraph = document.add_paragraph('', style='Body Text 2')
paragraph.add_run(
    '\n\nCONFIDENTIAL'
).bold = True
paragraph.add_run(
    '\n\n\nCPS Readiness Report'
).bold = True
paragraph.add_run(
    '\n\nOrganization: ' + organization_dictionary_aim[organization_number_aim]
).bold = True
paragraph.add_run(
    '\n\nDate Prepared: ' + current_date
).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_page_break()


# Table of Contents
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'Table of Contents\n\n'
).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'I.\tBackground Information\n\n'
)
paragraph.add_run(
    'II.\tQuantitative Assessment\n\n'
)
paragraph.add_run(
    'III.\tQualitative Assessment\n\n'
)
paragraph.add_run(
    'VI.\tCPS Readiness Summary\n\n'
)
paragraph.add_run(
    'V.\tRecommendations\n\n'
)
document.add_page_break()


# I. Background Information
paragraph = document.add_paragraph('\n\n', style='Normal')
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tI.\t'
)
paragraph.add_run(
    'Background Information'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tImplementation science has revealed that implementing any \
evidence-based approach requires changing the behavior of staff and the \
climate, culture, and structures of the organization. As a result, \
implementing with fidelity and sustainability requires time, patience, \
discomfort, perseverance, and good leadership.')
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tThe staff at your agency completed surveys to provide us with \
quantitative information. Additionally, focus groups with your agency’s \
staff gave us additional qualitative information that we have used toward \
this final report. These focus groups facilitated our understanding of \
current practices and challenges that your staff face. Every agency has \
challenges; understanding exactly what yours are will help us know how and \
when to implement CPS to maximize its success.'
)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tAfter you review this report, the Think:Kids team will discuss \
recommendations with organization leadership.  In some cases, immediate \
CPS implementation is recommended, and if so, a proposed comprehensive \
implementation and evaluation plan, timeline, and associated costs will be \
discussed.  In other cases, the Think:Kids team may recommend to leadership \
that particular readiness areas be addressed first, in order to maximize the \
benefits of implementation.'
)
document.add_page_break()

table = document.add_table(rows=1, cols=2, style='Normal Table')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Key:').bold = True
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Freq:').underline = True
row_cells[1].text = 'How many times a respondent responded in that way'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Percent:').underline = True
row_cells[1].text = 'The percent of the entire sample who responded in that \
way'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Cum:').underline = True
row_cells[1].text = 'Cumulative percent.  The percent of that response and \
all that came in list before'
row_cells = table.add_row().cells
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Obs:').underline = True
row_cells[1].text = 'Number of respondents who completed this item or for \
whom the subscale could be calculated in a valid way'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Mean:').underline = True
row_cells[1].text = 'Average of all relevant items'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Std Dev:').underline = True
row_cells[1].text = 'Standard deviation: Approximately 68% of responses fall \
within Mean +/- Standard deviation. For example, if the Mean is 3 and \
Standard deviation is 1, 68% of responses fall between 3+/-1, or between 2 \
and 4.'
row_cells = table.add_row().cells
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Min:').underline = True
row_cells[1].text = 'Minimum score for that item or category'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Max:').underline = True
row_cells[1].text = 'Maximum score for that item or category'


def set_col_widths(table, first_col, other_cols):
    widths = (Inches(first_col), Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


set_col_widths(table, 1, 6)

document.add_page_break()


# II. Quantitative Assessment
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tII.\t'
)
paragraph.add_run(
    'Quantitative Assessment'
).underline = True
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'The CPS Adherence and Impact Measure (CPS-AIM)'
).bold = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tThe CPS-AIM inquires about factors related to CPS. A pre-training \
measurement is taken so that we can see how staff’s adherence to the CPS \
philosophy increases over time. We also hope that over time we will see \
reduced burnout and perceptions of a more positive impact on youth.')
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Survey Responders:'
).bold = True
paragraph.add_run(
    '\n\t' + str(survey_count_aim) + ' total staff members responded to our CPS \
Adherence and Impact Measure Survey in a valid and reliable way. The \
quantitative analyses on the next few pages of this report are based on data \
collected from those respondents. The respondents are broken down by school, \
job role, and CPS training status as follows:\n'
)


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Staff Role Table
make_table(
    'Staff Role',
    role_count_aim,
    role_dictionary_aim,
    role_survey_count_aim
)

# Staff Program Table
if sub_school_dict:
    document.add_paragraph('', style='Normal')
    make_table(
        'Staff Program',
        sub_school_count,
        sub_school_dict,
        sub_school_survey_count
    )


# Staff Training Table
document.add_paragraph('', style='Normal')
table = document.add_table(rows=1, cols=4, style='Normal Table')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('CPS Training').bold = True
headers = ['Freq', 'Percent', 'Cum.']
for count, header in enumerate(headers):
    hdr_cells[count + 1].paragraphs[0].add_run(header).underline = True
total = 0
for key, value in training_dictionary.items():
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = str(value)
    total += int(value)
    row_cells[2].text = str(
        round(float((value / training_total) * 100), 0)) + '%'
    row_cells[3].text = str(
        round(float((total / training_total) * 100), 0)) + '%'
set_col_widths(table, 3, 1)


document.add_page_break()


# The CPS Adherence and Impact Measure (CPS-AIM)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tScores range from 1=Strongly Disagree to 7=Strongly Agree. \
In the graph below, the horizontal line indicates a score of 4, which is \
“Not Sure.” The goal is to be far above the horizontal line in Adherence to \
CPS Philosophy and Perceptions of Positive Impact, and far below the \
horizontal line in Burnout.'
)

# Creating the CPS AIM Educators Graph for Pre-Training
# Score the CPS AIM-E
cps_aim_columns = []
df_aim_columns(organization_data_aim, 'tkcote', cps_aim_columns)
cps_aim_columns.append('record_id')
cps_aim = organization_data_aim[cps_aim_columns]
cps_aim_columns_2 = []
for i in range(32):
    cps_aim_columns_2.append('tkcot_' + str(i + 1))
cps_aim_columns_2.append('record_id')
cps_aim.columns = cps_aim_columns_2
results = []
results = score.cps_aim_educator(cps_aim, results, 'record_id')

# Creating the CPS Adherence and Impact Measure Bar Graph
plt.figure()
results[results.columns[1:4]].mean().plot.bar(width=.9)
plt.suptitle('CPS Adherence and Impact Measure', fontsize=14)
plt.title('Pre-Training Score', fontsize=10)
plt.xticks((0, 1, 2),
           ('Adherence', 'Positive Impact', 'Burnout'),
           rotation=0)
plt.ylim(ymax=7, ymin=1)
plt.axhline(y=4, color='r', linestyle='-')
plt.ylabel('Mean Score')
plt.savefig('plt.png')
document.add_picture('plt.png')

# Adding the Summary: section, to be filled out later by Alisha
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Summary:'
).bold = True
paragraph.add_run(
    ' As a whole...'
)
document.add_page_break()

# Part 2!
# "CPS Readiness Assessment for Schools (1-Part)"
tk.readiness_for_schools()
project = Project(tk.api_url, tk.api_token)
data_readiness_educators = project.export_records(format='df')
metadata_readiness_educators = project.export_metadata(format='df')

organization_dictionary_readiness = meta_dict(
    metadata_readiness_educators, 'organization')

# Download the corresponding Readiness Survey Data
print("The organizations listed in the CPS Readiness Assessment for Schools \
(1-Part) are as follows.")
for key, value in organization_dictionary_readiness.items():
    print(str(key) + '\t' + value)
organization_number_readiness = int(input("Please enter the Organizations \
Number to generate a report on.\n\t"))

# Retrieving and sorting the Readiness Data from REDcap
data_readiness_educators.reset_index(inplace=True)
organization_data_readiness = data_readiness_educators[
    data_readiness_educators['organization'] == organization_number_readiness
]
# Remove participants who did not answer at least 15 questions
organization_data_readiness = organization_data_readiness.dropna(thresh=15)

# Replace any N/A answers (9) with np.nan so that the data is not skewed
organization_data_readiness = organization_data_readiness.replace(9, np.nan)

# Reverse the Reverse scored columns
reverse = [
    'staff5', 'staff13', 'admin5', 'admin13'
]
for i in reverse:
    organization_data_readiness[i].replace(
        [1, 2, 3, 4, 5], [5, 4, 3, 2, 1],
        inplace=True)
# Variables from Readiness Data Set

# Get a count of how many participants responded to the survey and answered
# at least 15 questions.
survey_count_readiness = organization_data_readiness.record_id.count()

# Role variables
role_dictionary_readiness = meta_dict(
    metadata_readiness_educators, 'rolestaff')
role_count_readiness = (
    organization_data_readiness['rolestaff'].value_counts().to_dict()
)
role_survey_count_readiness = organization_data_readiness['rolestaff'].count()

# Years at Organization Dictionary - Readiness
years_at_org = meta_dict(metadata_readiness_educators, 'yearsatorg')
years_at_org_count = (
    organization_data_readiness['yearsatorg'].value_counts().to_dict())
years_at_org_total_count = organization_data_readiness['yearsatorg'].count()
# Training Dictionary
training_dictionary = meta_dict(metadata_readiness_educators, 'training')
training_count = (
    organization_data_readiness['training'].value_counts().to_dict())
training_total_count = organization_data_readiness['training'].count()

# Sub School
sub_school_column = []
for x, i in enumerate(metadata_readiness_educators['branching_logic']):
    if str(organization_number_readiness) in str(i):
        sub_school_column.append(metadata_readiness_educators.index[x])
sub_school_dict = {}

for i in sub_school_column:
    if ' |' in i:
        sub_school_dict = meta_dict(metadata_readiness_educators, i)
        sub_school_count = organization_data_readiness[i].value_counts(
        ).to_dict()
        sub_school_survey_count = organization_data_readiness[i].count()

# CPS Readiness Assessment
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'The CPS Readiness Survey'
).bold = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tEvaluating readiness to implement an evidence-informed \
approach like CPS revolves around several factors. These include an agency’s ')
paragraph.add_run(
    'motivation for change, its general capacity'
).italic = True
paragraph.add_run(
    ' for implementation of any intervention, and its '
)
paragraph.add_run(
    'specific capacity for implementation of the intervention in question'
).italic = True
paragraph.add_run(
    ' (in this case, CPS). These factors can vary by program within the \
agency.'
)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tUnder the category of '
)
paragraph.add_run(
    'Motivation for Change'
).bold = True
paragraph.add_run(
    ', we assess whether the administrators/leaders as well as other staff \
see the need, and have enthusiasm for, a different or additional approach. ')
paragraph.add_run(
    'Capacity in General'
).bold = True
paragraph.add_run(
    ' refers to things such as whether staff feel appropriately supported and \
satisfied with their work, whether sufficient supervision, communication and \
documentation structures are in place, and whether there is strong leadership \
present to facilitate implementation. '
)
paragraph.add_run(
    'Capacity for CPS'
).bold = True
paragraph.add_run(
    ' refers to a site’s ability to implement CPS in particular. For example, \
because CPS typically requires a significant shift in mindset, culture, and \
behavior from staff, it is even more important that influential \
culture-carriers be present. In addition, existing models of intervention \
must not directly conflict with the basic tenets of CPS that run counter to \
many conventional approaches aimed at motivating more compliant behavior \
externally through use of rewards and punishments.'
)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tTo conduct a comprehensive readiness assessment across programs, we \
utilize a readiness measure designed explicitly for this purpose and based on \
the latest research on organizational readiness for implementation of an \
innovation (Scaccia et al., 2015). ')
document.add_page_break()

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Survey Responders:'
).bold = True
paragraph.add_run(
    '\n\t' + str(survey_count_readiness) + ' total staff members responded to \
our CPS Readiness Surveys in a valid and reliable way. The quantitative \
analyses on the next few pages of this report are based on data collected \
from those respondents. The respondents are broken down by school, job role, \
years of employment, and CPS training status as follows:'
)

# Staff Role Table Readiness
make_table(
    'Staff Role',
    role_count_readiness,
    role_dictionary_readiness,
    role_survey_count_readiness
)

# Staff Program Table Readiness

if sub_school_dict:
    document.add_paragraph('', style='Normal')
    make_table(
        'Staff Program',
        sub_school_count,
        sub_school_dict,
        sub_school_survey_count
    )

# Staff Training Table Readiness
document.add_paragraph('', style='Normal')
make_table(
    'Training in CPS',
    training_count,
    training_dictionary,
    training_total_count
)

# Years at Organization Table Readiness
document.add_paragraph('', style='Normal')

make_table(
    'Years at Organization',
    years_at_org_count,
    years_at_org,
    years_at_org_total_count
)

document.add_page_break()

# Readiness Survey Results, for All Employees
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Readiness Survey Results, for All Employees'
).bold = True

# Score the Readiness Survey
# Score the Educators Version
results_educators = []
results_educators = score.cps_readiness_educator(
    organization_data_readiness,
    results_educators,
    'record_id')
results_educators = results_educators.dropna(thresh=3)

# Score the Admin Version
results_admin = []
results_admin = score.cps_readiness_admin(
    organization_data_readiness,
    results_admin,
    'record_id')
results_admin = results_admin.dropna(thresh=3)

frames = [results_educators, results_admin]
results = pd.concat(frames)

# Add the scored columns back into the main DataFrame.
organization_data_readiness = pd.merge(
    left=organization_data_readiness,
    right=results,
    left_on='record_id',
    right_on='id')

# CPS Readiness Measure Score Variables.


def stats_dict(df, column):
    dictionary = {'Count': df[column].count(),
                  'Mean': round(df[column].mean(), 1),
                  'Standard Deviation': round(df[column].std(), 2),
                  'Minimum': round(df[column].min(), 2),
                  'Maximum': round(df[column].max(), 2)
                  }
    return dictionary


# BIG DICTIONARY OF READINESS INFORMATION!!!
readiness_data = {
    'Motivation': {
        'All Staff': {
            'Overall Stats': stats_dict(
                organization_data_readiness, 'readiness_motiv_mean'),
        },
        'Administration': {
            'Overall': stats_dict(
                results_admin, 'readiness_motiv_mean'),
            'admin1': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin1',
                    'field_label'),
                'Trunc Question': 'Policies need improvement',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin1')
            },
            'admin2': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin2',
                    'field_label'),
                'Trunc Question': 'CPS is improvement',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin2')
            },
            'admin5': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin5',
                    'field_label'),
                'Trunc Question': 'CPS too hard (reversed)',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin5')
            },
            'admin6': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin6',
                    'field_label'),
                'Trunc Question': 'Leaders want all in',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin6')
            },
            'admin7': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin7',
                    'field_label'),
                'Trunc Question': 'CPS consistent with values',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin7')
            },
            'admin8': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin8',
                    'field_label'),
                'Trunc Question': 'CPS consistent with practice',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin8')
            },
        },
        'Staff': {
            'Overall': stats_dict(
                results_educators, 'readiness_motiv_mean'),
            'staff1': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff1',
                    'field_label'),
                'Trunc Question': 'Policies need improvement',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff1')
            },
            'staff2': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff2',
                    'field_label'),
                'Trunc Question': 'CPS is improvement',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff2')
            },
            'staff5': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff5',
                    'field_label'),
                'Trunc Question': 'CPS too hard (reversed)',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff5')
            },
            'staff6': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff6',
                    'field_label'),
                'Trunc Question': 'Leaders want all in',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff6')
            },
            'staff7': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff7',
                    'field_label'),
                'Trunc Question': 'CPS consistent with values',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff7')
            },
            'staff8': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff8',
                    'field_label'),
                'Trunc Question': 'CPS consistent with practice',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff8')
            },
        }
    },
    'General Capacity': {
        'All Staff': {
            'Overall Stats': stats_dict(
                organization_data_readiness, 'readiness_capacity_mean'),
        },
        'Administration': {
            'Overall': stats_dict(
                results_admin, 'readiness_capacity_mean'),
            'admin9': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin9',
                    'field_label'),
                'Trunc Question': 'We encourage innovation',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin9')
            },
            'admin10': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin10',
                    'field_label'),
                'Trunc Question': 'Staff want to learn more',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin10')
            },
            'admin11': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin11',
                    'field_label'),
                'Trunc Question': 'Staff supported by leaders',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin11')
            },
            'admin12': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin12',
                    'field_label'),
                'Trunc Question': 'Staff communicate well',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin12')
            },
            'admin13': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin13',
                    'field_label'),
                'Trunc Question': 'Too many interventions (reversed)',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin13')
            },
        },
        'Staff': {
            'Overall': stats_dict(
                results_educators, 'readiness_capacity_mean'),
            'staff9': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff9',
                    'field_label'),
                'Trunc Question': 'We encourage innovation',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff9')
            },
            'staff10': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff10',
                    'field_label'),
                'Trunc Question': 'Staff want to learn more',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff10')
            },
            'staff11': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff11',
                    'field_label'),
                'Trunc Question': 'Staff supported by leaders',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff11')
            },
            'staff12': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff12',
                    'field_label'),
                'Trunc Question': 'Staff communicate well',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff12')
            },
            'staff13': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff13',
                    'field_label'),
                'Trunc Question': 'Too many interventions (reversed)',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff13')
            },
        }
    },
    'CPS Capacity': {
        'All Staff': {
            'Overall Stats': stats_dict(
                organization_data_readiness, 'readiness_cps_cap_mean'),
        },
        'Administration': {
            'Overall': stats_dict(
                results_admin, 'readiness_cps_cap_mean'),
            'admin3': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin3',
                    'field_label'),
                'Trunc Question': 'Leaders are committed',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin3')
            },
            'admin4': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin4',
                    'field_label'),
                'Trunc Question': 'Internal CPS team',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin4')
            },
            'admin14': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'admin14',
                    'field_label'),
                'Trunc Question': 'Financially committed',
                'Stats': stats_dict(
                    organization_data_readiness, 'admin14')
            }
        },
        'Staff': {
            'Overall': stats_dict(
                results_educators, 'readiness_cps_cap_mean'),
            'staff3': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff3',
                    'field_label'),
                'Trunc Question': 'Leaders are committed',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff3')
            },
            'staff4': {
                'Full Question': meta_cell_content(
                    metadata_readiness_educators,
                    'staff4',
                    'field_label'),
                'Trunc Question': 'Internal CPS team',
                'Stats': stats_dict(
                    organization_data_readiness, 'staff4')
            }
        }
    }
}

# Admin
admin_readiness_results = {
    'Motivation for Change': stats_dict(
        results_admin, 'readiness_motiv_mean'),
    'General Capacity': stats_dict(
        results_admin, 'readiness_capacity_mean'),
    'Capacity for CPS': stats_dict(
        results_admin, 'readiness_cps_cap_mean')
}

# Educators
educators_readiness_results = {
    'Motivation for Change': stats_dict(
        results_educators, 'readiness_motiv_mean'),
    'General Capacity': stats_dict(
        results_educators, 'readiness_capacity_mean'),
    'Capacity for CPS': stats_dict(
        results_educators, 'readiness_cps_cap_mean'),
}

# All
all_readiness_results = {
    'Motivation for Change': stats_dict(
        organization_data_readiness, 'readiness_motiv_mean'),
    'General Capacity': stats_dict(
        organization_data_readiness, 'readiness_capacity_mean'),
    'Capacity for CPS': stats_dict(
        organization_data_readiness, 'readiness_cps_cap_mean')
}


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def readiness_tables(dictionary):
    table = document.add_table(rows=1, cols=6, style='Normal Table')
    hdr_cells = table.rows[0].cells
    headers = ['Variable:', 'Obs', 'Mean', 'Std. Dev.', 'Min', 'Max']
    for count, header in enumerate(headers):
        hdr_cells[count].paragraphs[0].add_run(header).underline = True
    for key, value in dictionary.items():
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(key)
        for x, y in enumerate(value.values()):
            row_cells[x + 1].text = str(y)
    set_col_widths(table, 2, 1)


# Tables describing statistics for all staff
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tPossible responses for readiness items range from 1 (Strongly Disagree)\
 to 5 (Strongly Agree), with a 3 response for "Not Sure."')
readiness_tables(all_readiness_results)
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Summary:  Overall, staff at \
' + organization_dictionary_aim[organization_number_aim] + ' are…\
\n\nThis spread can be seen in more detail in the histograms below.')

document.add_page_break()

# Create the Histrogram for The CPS Readiness Survey
# Motivation, General Capacity, and Capacity for CPS, for ALL STAFF
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Detail View:  Motivation, General Capacity, and Capacity for CPS, \
for ALL STAFF'
).bold = True


def all_staff_readiness_plots(title, column):
    plt.figure(figsize=(3, 3))
    organization_data_readiness[column].plot.hist(bins=np.linspace(1, 5, 9))
    plt.title(title, fontsize=14)
    plt.xticks((1, 2, 3, 4, 5))
    plt.tight_layout()
    plt.savefig('plt.png')
    paragraph.add_run().add_picture('plt.png')


all_staff_readiness_plots(
    'Motivation of All Staff',
    'readiness_motiv_mean')

all_staff_readiness_plots(
    'General Capacity of All Staff',
    'readiness_capacity_mean')

all_staff_readiness_plots(
    'CPS Capacity of All Staff',
    'readiness_cps_cap_mean')


document.add_page_break()

# *** Alisha might want to combine these two tables? ***
# Tables describing statistics for Staff
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Readiness Survey Results Staff'
).bold = True
readiness_tables(educators_readiness_results)

# Tables describing statistics for Administration
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run(
    'Readiness Survey Results Leadership/Administration'
).bold = True
readiness_tables(admin_readiness_results)

paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Summary: Across roles, staff at are…')

document.add_page_break()

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Detail View: Motivation, General Capacity, and Capacity for CPS, by Role:'
).bold = True

role_count_threshold = int(
    input('Under what number will we exclude staff from the analysis?\n\t')
)
count_role_graphs = []
for x, i in enumerate(organization_data_readiness['rolestaff']):
    if role_count_readiness[i] < role_count_threshold:
        organization_data_readiness.drop(x, inplace=True)
    else:
        count_role_graphs.append(i)


count_role_graphs = int(np.ceil(len(set(count_role_graphs)) / 2))
count_role_graphs = int(count_role_graphs * 2)

unique_rolestaff = organization_data_readiness['rolestaff'].nunique()
unique_rolestaff = int(np.ceil(unique_rolestaff / 2))
organization_data_readiness[
    'rolestaff'
] = organization_data_readiness['rolestaff'].map(
    role_dictionary_readiness)


paragraph = document.add_paragraph('\t', style='Normal')
paragraph.add_run(
    'In order to preserve confidentiality, categories where there were less \
than ' + str(role_count_threshold) + ' \
participants who answered the survey in their role were excluded from the \
following analyses.')


def by_role_hist(title, column):
    paragraph = document.add_paragraph('', style='Body Text')
    paragraph.add_run(title).underline = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.figure()
    organization_data_readiness[column].hist(
        figsize=(6.4, count_role_graphs),
        by=organization_data_readiness['rolestaff'],
        bins=np.linspace(1, 5, 9),
        layout=(unique_rolestaff, 2))
    plt.xticks((1, 2, 3, 4, 5))
    plt.tight_layout()
    plt.savefig('plt.png')
    paragraph.add_run().add_picture('plt.png')
    document.add_page_break()


by_role_hist('Motivation by Role', 'readiness_motiv_mean')
by_role_hist('General Capacity by Role', 'readiness_capacity_mean')
by_role_hist('CPS Capacity by Role', 'readiness_cps_cap_mean')

# ADD STAFF AND ADMIN TABLES
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Detailed Report of ').bold = True
p = paragraph.add_run('Educational Staffs’')
p.bold = True
p.underline = True
paragraph.add_run(' Responses by Item:').bold = True


def detailed_report(kind):
    table = document.add_table(rows=1, cols=6, style='Normal Table')
    hdr_cells = table.rows[0].cells
    headers = ['Question:', 'Obs', 'Mean', 'Std. Dev.', 'Min', 'Max']
    for count, header in enumerate(headers):
        hdr_cells[count].paragraphs[0].add_run(header).underline = True
    for key, value in readiness_data.items():
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(key).bold = True
        for key, value in value.items():
            if kind.title() in key:
                for key, value in value.items():
                    if kind.lower() in key:
                        question_number = key[5:]
                        for key, value in value.items():
                            if key == 'Trunc Question':
                                row_cells = table.add_row().cells
                                row_cells[0].paragraphs[0].add_run(
                                    str(question_number) + ') ' + value)
                            if key == 'Stats':
                                for x, value in enumerate(value.values()):
                                    row_cells[x + 1].text = str(value)
    set_col_widths(table, 2.9, .9)


detailed_report('staff')

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\t* Items have been truncated here to conserve space.\nSee Appendix for \
item wording.'
).bold = True
document.add_page_break()

paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Detailed Report of ').bold = True
p = paragraph.add_run('Leadership/Administrators’')
p.bold = True
p.underline = True
paragraph.add_run(' Responses by Item:').bold = True

detailed_report('admin')

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\t* Items have been truncated here to conserve space.\nSee Appendix for \
item wording.'
).bold = True
document.add_page_break()

# III. Qualitative Assessment
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tIII.\t'
)
paragraph.add_run(
    'Qualitative Assessment'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Summary of Need:'
).bold = True
paragraph.add_run('\n')
paragraph.add_run(
    '\nReadiness strengths:'
).bold = True
paragraph.add_run('\n')
paragraph.add_run(
    '\nReadiness areas in need of improvement:'
).bold = True
paragraph.add_run('\n')
paragraph = document.add_paragraph(
    '\nReadiness Area: Motivation for Change', style='Normal'
)
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph(
    '\nReadiness Area: Capacity in General', style='Normal'
)
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph(
    '\nReadiness Area: Specific Capacity for CPS', style='Normal'
)
document.add_paragraph(
    '', style='List Bullet 2'
)
document.add_page_break()


# IV. CPS Readiness Summary
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tIV.\t'
)
paragraph.add_run(
    'CPS Readiness Summary'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run('Insert summary rating form after quantitative interviews')

readiness_columns_admin = []
readiness_columns_admin_numbers = []
readiness_columns_staff = []
readiness_columns_staff_numbers = []

for key, value in readiness_data.items():
    for key, value in value.items():
        if key == 'Staff' or key == 'Administration':
            for key, value in value.items():
                if 'admin' in key:
                    readiness_columns_admin.append(key)
                    readiness_columns_admin_numbers.append(int(key[5:]))
                elif 'staff' in key:
                    readiness_columns_staff.append(key)
                    readiness_columns_staff_numbers.append(int(key[5:]))

readiness_columns_admin_numbers.sort()
readiness_columns_staff_numbers.sort()
admin = organization_data_readiness[readiness_columns_admin].dropna(thresh=13)
admin = admin.replace(np.nan, 0)
admin.columns = readiness_columns_admin_numbers
admin = admin.reset_index().drop(['index'], axis=1).astype(int)
staff = organization_data_readiness[readiness_columns_staff].dropna(thresh=10)
staff = staff.replace(np.nan, 0)
staff.columns = readiness_columns_staff_numbers
staff = staff.reset_index().drop(['index'], axis=1).astype(int)

# Can get around 30 participants on each heatmap. Make function that will loop
# Through the dataframe and make a heatmap for each 40 participant groups


def heat_map(df, group):
    height = (df[1].count() * .18) + 1
    sns.set()
    f, ax = plt.subplots(figsize=(6.5, height))
    sns.heatmap(
        df,
        annot=True,
        fmt="d",
        ax=ax,
        robust=True,
        cbar=False,
        vmin=1,
        vmax=5,
        cmap="YlGnBu")
    plt.title('Readiness Survey Responses Heat Map for ' + group)
    plt.xlabel('Question Number')
    plt.ylabel('Participants')
    plt.tight_layout()
    plt.savefig('plt.png')
    paragraph.add_run().add_picture('plt.png')


x = []
for i in staff.index:
    x.append(i)
    if i == 0:
        continue
    if i % 40 == 0:
        heat_map(staff.loc[x], 'Staff')
        x = []
if len(x) < 41:
    heat_map(staff.loc[x], 'Staff')

x = []
for i in admin.index:
    x.append(i)
    if i == 0:
        continue
    if i % 40 == 0:
        heat_map(admin.loc[x], 'Admin')
        x = []
if len(x) < 41:
    heat_map(admin.loc[x], 'Admin')

document.add_page_break()

# V. Recommendations
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tV.\t'
)
paragraph.add_run(
    'Recommendations'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tBased upon the complete results of this readiness assessment, our \
Readiness Team feels \
that ' + organization_dictionary_aim[organization_number_aim] + ' is in \
excellent shape to continue to the next phase of implementation. Our \
recommendations are as follows:'
)
paragraph.add_run('\n')

paragraph = document.add_paragraph(
    '\nPrior to Training', style='Normal'
).bold = True
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph(
    '\nTraining and Coaching', style='Normal'
).bold = True
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph(
    '\nMoving Toward Sustainability', style='Normal'
).bold = True
document.add_paragraph(
    '', style='List Bullet 2'
)
document.add_page_break()

# Appendix
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'Appendix'
).underline = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'School/Organization-Wide Implementation Readiness Survey'
).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\nFOR STAFF:'
)

question_columns_staff = []

questions_staff = []
for i in metadata_readiness_educators.index:
    if 'staff' in i:
        question_columns_staff.append(i)
        question = meta_cell_content(
            metadata_readiness_educators, i, 'field_label'
        )
        for i in question:
            questions_staff.append(i)
questions_staff = questions_staff[1:]
question_columns_staff = question_columns_staff[1:]


def add_survey(job_type):
    count = 0
    for i in job_type:
        count += 1
        paragraph = document.add_paragraph(str(count) + ') ', style='Normal')
        paragraph.add_run(i)
        if count in [1, 2, 5, 6, 7, 8]:
            paragraph.add_run(' (Motivation for Change)')
        elif count in [3, 4]:
            paragraph.add_run(' (CPS Capacity)')
        else:
            paragraph.add_run(' (General Capacity)')
        if count in [5, 13]:
            paragraph.add_run(' (Reversed)')


add_survey(questions_staff)
document.add_page_break()
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\nFOR LEADERS/ADMINISTRATION:'
)

question_columns_admin = []
questions_admin = []
for i in metadata_readiness_educators.index:
    if 'admin' in i:
        question_columns_admin.append(i)
        question = meta_cell_content(
            metadata_readiness_educators, i, 'field_label'
        )
        for i in question:
            questions_admin.append(i)
add_survey(questions_admin)

document.save(organization_dictionary_aim[organization_number_aim] + ' \
Readiness Report.docx')
