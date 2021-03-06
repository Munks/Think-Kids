from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import datetime as dt
from redcap import Project
import os
import sys
sys.path.append(r'C:\Python Programs\Think-Kids_Private')
import tokens as tk
import scoring as score


"""
Code to create a Readiness Report. The data is pulled from REDcap. Data is
retrieved any of the following three locations:

    "CPS-AIM, Educators' Version"
        "CPS Readiness Assessment for Schools (1-Part)"

    "CPS-AIM, Systems' Version"
        "CPS Readiness Assessment for Agencies (1-Part)"

"""


# Functions
def add_document_text_from_excel(df):
    """
    This takes text out of a formatted excel and adds it to the document.
    """
    for i in df.index:
        if df.iloc[i]['Paragraph']:
            if df.iloc[i]['Paragraph'] == 'Skip':
                continue
            else:
                paragraph_style = str(df.iloc[i]['Paragraph'])
        if df.iloc[i]['Text']:
            text = str(df.iloc[i]['Text'])
        if df.iloc[i]['Style']:
            text_style = str(df.iloc[i]['Style'])
        if 'nan' not in paragraph_style:
            paragraph = document.add_paragraph('', style=paragraph_style)
        if text_style == 'bold':
            paragraph.add_run(text).bold = True
        elif text_style == 'italic':
            paragraph.add_run(text).italic = True
        elif text_style == 'underline':
            paragraph.add_run(text).underline = True
        else:
            paragraph.add_run(text)


def add_survey(job_type):
    """
    This adds the appendix surveys to the end of the Word Document
    """
    count = 0
    for i in job_type:
        count += 1
        paragraph = document.add_paragraph(str(count) + ') ', style='Normal')
        paragraph.add_run(i)
        if count in [1, 2, 5, 6, 7, 8]:
            paragraph.add_run(' (Fit)')
        else:
            paragraph.add_run(' (Capacity)')
        if count in [5, 13]:
            paragraph.add_run(' (Reversed)')


def cell_content(df, row, column):
    """
    This function saves the content of a single cell in a DataFrame as a list
    """
    content = df.loc[[row], [column]]
    content = content[column]
    content_list = []
    for i in content:
        content_list.append(i)
    return content_list


def meta_dict(df, row):
    """
    Create a python dictionary from the Data Dictionary downloaded from REDcap

    df = Data Dictionary DataFrame Pandas Object.
    row = name of row in the Data Dictionary where you are looking for
        dictionary keys and values.
    The column name in the Data Dictionary is with keys and values is a
        constant and always 'select_choices_or_calculations'.
    """
    column = 'select_choices_or_calculations'
    content = cell_content(df, row, column)
    dictionary_keys = []
    dictionary_values = []
    for i in content:
        content = i.split(' | ')
        # Divide the string into lines as displayed in REDcap.
        for i in content:
            content = i.split(', ', 1)
            # Divide the lines into Answer and Key as input in REDcap and save
            #   as a list.
            a = 0
            for i in content:
                # First item in the list is the key. Add to the Keys List.
                if a == 0:
                    dictionary_keys.append(int(i))
                    a += 1
                # Second item in the list is the value. Add to the Values List.
                else:
                    dictionary_values.append(i)
    dictionary = dict(zip(dictionary_keys, dictionary_values))
    return dictionary


def font_style(style_name, size, font_name):
    """
    style_name = name of the style as it appears in the Word Doc Template
    size = font size
    font_name = name of font in Word
    """
    font = document.styles[style_name].font
    font.name = font_name
    font.size = Pt(size)


def make_table(title, primary, secondary, count):
    """
    Creates Readiness Report Tables

    title = The title of the table as it will appear in the Word Document.
    primary = Dictionary with numeric value of the variable and count of how
    #   many times that variable appears.
    secondary = Dictionary that links the numeric value of the variable with
    #   the name as it will appear in the table.
    count = Count of total of how many times all variables appear.
    """
    table = document.add_table(rows=1, cols=3, style='Normal Table')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run(title).bold = True
    headers = ['Freq', 'Percent']
    for position, header in enumerate(headers):
        hdr_cells[position + 1].paragraphs[0].add_run(header).underline = True
    total = 0
    for key, value in primary.items():
        row_cells = table.add_row().cells
        row_cells[0].text = secondary[key]
        row_cells[1].text = str(value)
        total += int(value)
        row_cells[2].text = str(round(float(value / count) * 100, 0)) + '%'
    set_col_widths(table, 4, 1)


def stats_dict(df, column):
    """
    Creates a dictionary of classic discriptive stats based on column data from
        the DataFrame.

    df = DataFrame
    column = Column of data to calculate stats for
    """
    dictionary = {'Count': df[column].count(),
                  'Mean': round(df[column].mean(), 1),
                  'Standard Deviation': round(df[column].std(), 2),
                  'Minimum': round(df[column].min(), 2),
                  'Maximum': round(df[column].max(), 2)
                  }
    return dictionary


def readiness_tables(dictionary):
    """
    Creates the Readiness Tables in the Word Doc based on the given dictionary
    """
    table = document.add_table(rows=1, cols=6, style='Normal Table')
    hdr_cells = table.rows[0].cells
    headers = ['Variable:', 'Obs', 'Mean', 'Std. Dev.', 'Min', 'Max']
    for count, header in enumerate(headers):
        hdr_cells[count].paragraphs[0].add_run(header).underline = True
    for key, value in dictionary.items():
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(key)
        for x, y in enumerate(value.values()):
            row_cells[x + 1].text = str(y)
    set_col_widths(table, 2, 1)


def histogram_general(title, df, column):
    """
    Creates a single histogram of the column data

    title = Title of the Graph as it will appear in the Word Doc
    df = location of the data
    column = column in the data frame with the required data
    """
    plt.figure(figsize=(3, 3))
    df[column].plot.hist(bins=np.linspace(1, 5, 9))
    plt.title(title, fontsize=14)
    plt.xticks((1, 2, 3, 4, 5))
    plt.tight_layout()
    plt.savefig('plt.png')
    paragraph.add_run().add_picture('plt.png')


def histogram_by(title, df, df_column, sort_by, height, columns):
    """
    Creates multiple histograms of the column data broken down by role

    title = Title of the Graph as it will appear in the Word Doc
    df = location of the data
    column = column in the data frame with the required data
    sort_by = the column used to catagorize the data
    """
    paragraph = document.add_paragraph('', style='Body Text')
    paragraph.add_run(title).underline = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.figure()
    df[df_column].hist(
        figsize=(6.4, height),
        by=df[sort_by],
        bins=np.linspace(1, 5, 9),
        layout=(columns, 2),
        sharey=True,
        sharex=True,
        xrot=90)
    plt.xticks((1, 2, 3, 4, 5))
    plt.tight_layout()
    plt.savefig('plt.png')
    paragraph.add_run().add_picture('plt.png')
    document.add_page_break()


def detailed_report(kind):
    table = document.add_table(rows=1, cols=6, style='Normal Table')
    hdr_cells = table.rows[0].cells
    headers = ['Question:', 'Obs', 'Mean', 'Std. Dev.', 'Min', 'Max']
    for count, header in enumerate(headers):
        hdr_cells[count].paragraphs[0].add_run(header).underline = True
    for key, value in readiness_data.items():
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(key).bold = True
        for key, value in value.items():
            if kind.title() in key:
                for key, value in value.items():
                    if kind.lower() in key:
                        question_number = key[5:]
                        for key, value in value.items():
                            if key == 'Trunc Question':
                                row_cells = table.add_row().cells
                                row_cells[0].paragraphs[0].add_run(
                                    str(question_number) + ') ' + value)
                            if key == 'Stats':
                                for x, value in enumerate(value.values()):
                                    row_cells[x + 1].text = str(value)
    set_col_widths(table, 2.9, .9)


def heat_map(df, chart_title, min, max, height, width):
    """
    Create the Heat Map of our data. This shows how each participant answered
        on the readiness measure
    """
    sns.set()
    f, ax = plt.subplots(figsize=(width, height))
    sns.heatmap(
        df,
        annot=False,
        ax=ax,
        robust=True,
        cbar=False,
        vmin=min,
        vmax=max,
        cmap="YlGnBu",
        yticklabels=False)
    plt.title(chart_title)
    plt.xticks(rotation=90)
    plt.ylabel('Participants')
    plt.tight_layout()
    plt.savefig('plt.png')


# Import the Readiness Source Text

source_text = pd.read_excel(
    'readiness_source_text.xlsx',
    sheet_name=None)

readiness_survey_text = source_text['CPS Readiness Survey']
cps_aim_subscales_text = source_text['CPS-AIM Subscales']
total_readiness_table_text = source_text['Total Readiness Table']

# User Interface
#   Think:Kids
print('\n\n')
print("\t\t* * * * * * *   *    *      *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *  *")
print("\t\t      *              **")
print("\t\t      *         *    *  *")
print("\t\t      *        * *   *    *")
print("\t\t      *         *    *      *")
print("\n\t\t\tReadiness Reports!!\n\n")

tk.cps_aim_pre_readiness()


# Global Variables
#   In the CPS-AIM, one survey has the variable for the question:
#       'What group or organization asked you to fill out this form?'
#   as organization, the other as org, this variable differentiates between the
#   two options.
cps_aim_org = 'org'
#   In the Readiness Assessment this variable is 'organization' for both
#   surveys
cps_readiness_org = 'organization'
#   Current Date
current_date = str(dt.date.today().strftime('%m-%d-%Y'))


# Download the CPS-AIM off of REDcap
project = Project(tk.api_url, tk.api_token)
# CPS-AIM Data
cps_aim_df = project.export_records(format='df')
# CPS-AIM MetaData
meta_df = project.export_metadata(format='df')

organization_dictionary_aim = meta_dict(meta_df, cps_aim_org)

# Determine if the report is going to be on a system, or a school
choice_dict = {
    1: 'school',
    2: 'system',
}

# User Interface
#   Ask the User what organization they are generating a report for.
print("The organizations listed in the CPS-AIM, Longitudinal\
 Version are as follows.")
for key, value in organization_dictionary_aim.items():
    print(str(key) + '\t' + value)
cps_aim_choice = int(input("Please enter the Organizations Number to \
generate a report on.\n\t"))
print('Thanks!\n\n\n')

#   Ask the User if they are generating a report on systems, or schools
print('Is this organization a school or a system?\n')
for key, value in choice_dict.items():
    print(str(key) + '\t' + value.title())
choice = int(input('\n\t'))
print('Thanks!\n\n\n')

# Set the choice to be either a school or system.
choice = choice_dict[choice]

# Save the organization name
org = organization_dictionary_aim[cps_aim_choice]

# Retrieving and sorting the AIM Data from REDcap
cps_aim_df.reset_index(inplace=True)
cps_aim_df = cps_aim_df[cps_aim_df[cps_aim_org] == cps_aim_choice]

# Replace any N/A answers (9) with np.nan so that the data is not skewed
cps_aim_df = cps_aim_df.replace(9, np.nan)

# Remove participants who did not answer at least 30 questions
cps_aim_df.dropna(thresh=30, inplace=True)

# Sub School
# Only if there is a Sub School Question. Every school moving forward will have
# a Sub School Question
sub_org_variable_name = []
for x, i in enumerate(meta_df['branching_logic']):
    # Check the branching logic field in the Data Dictionary to see if the
    # Organization we entered is present
    if '[' + cps_aim_org + '] = ' + "'" + str(
            cps_aim_choice) + "'" in str(i):
        sub_org_variable_name.append(meta_df.index[x])

sub_org_dict = {}
for i in sub_org_variable_name:
    sub_org_variable_name = i
    sub_org_dict = meta_dict(meta_df, i)
    sub_org_count = cps_aim_df[i].value_counts().to_dict()
    sub_org_survey_count = cps_aim_df[i].count()

# User Interface
#   Ask the user if they want to break the data down by Sub-Catagory.
breakdown = int(input("Generate a report for:\
\n1\tThe organization as a whole\
\n2\tA specific program within the Organization\n\t"))
print('\nThanks!\n\n\n')
# If 2, ask the user which sub-catagory they would like to generate a report
# for. If 1, the code will skip this.
if breakdown == 2:
    for key, value in sub_org_dict.items():
        print(str(key) + '\t' + value)
    cps_aim_choice = int(input("Please enter the program \
to generate a report on.\n\t"))
    print('Thanks!\n\n\n')
    cps_aim_df = (
        cps_aim_df[cps_aim_df[sub_org_variable_name] == cps_aim_choice])
    org = str(org) + ', ' + str(sub_org_dict[cps_aim_choice])

# Variables from AIM Data Set

# Role variables educators
role_dictionary_aime = meta_dict(meta_df, 'rolee')
role_count_aime = cps_aim_df['rolee'].value_counts().to_dict()
role_survey_count_aime = cps_aim_df['rolee'].count()

# Role variables clinical
role_dictionary_aimc = meta_dict(meta_df, 'rolec')
role_count_aimc = cps_aim_df['rolec'].value_counts().to_dict()
role_survey_count_aimc = cps_aim_df['rolec'].count()
# Get a count of how many participants responded to the survey and answered
# at least 20 questions.
survey_count_aim = cps_aim_df.record_id.count()

# Create a dataframe of just the training columns
training = cps_aim_df.training

# Sort the information into a dictionary
training_dictionary = {
    'No Training': 0,
    'Not Sure': 0,
    'Received Training': 0,
}

# Build a count of the total number of people who answered anything about
# training
training_total = 0
for i in training:
    # An answer of 0 means they received no training
    if i == 2:
        training_dictionary['No Training'] += 1
        training_total += 1
    # An answer of 2 means they are not sure
    if i == 200:
        training_dictionary['Not Sure'] += 1
        training_total += 1
    if i == 1:
        training_dictionary['Received Training'] += 1
        training_total += 1


# Creating the Report in a Word Docutment
# Any code with "paragraph", "document", or "table" is responsible for setting
# up the final word document.
document = Document('Readiness Report Template.docx')
# Setting up the document fonts and styles
styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]

font_style('Normal', 12, 'Times New Roman')
font_style('No Spacing', 12, 'Times New Roman')
font_style('Body Text', 14, 'Times New Roman')
font_style('Body Text 2', 18, 'Times New Roman')
font_style('Title', 16, 'Times New Roman')


# Title Page
paragraph = document.add_paragraph('', style='Body Text 2')
paragraph.add_run('\n\nCONFIDENTIAL').bold = True
paragraph.add_run('\n\n\nCPS Readiness Report').bold = True
paragraph.add_run('\n\n' + choice.title() + ': ' + org).bold = True
paragraph.add_run('\n\nDate Prepared: ' + current_date).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_page_break()


# Table of Contents
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'Table of Contents\n\n'
).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'I.\tBackground Information\n\n'
)
paragraph.add_run(
    'II.\tQuantitative Assessment\n\n'
)
paragraph.add_run(
    'III.\tQualitative Assessment\n\n'
)
paragraph.add_run(
    'IV.\tCPS Readiness Summary\n\n'
)
paragraph.add_run(
    'V.\tRecommendations\n\n'
)
document.add_page_break()


# I. Background Information
paragraph = document.add_paragraph('\n\n', style='Normal')
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tI.\t'
)
paragraph.add_run(
    'Background Information'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tImplementation Science has revealed that implementing any \
evidence-based approach requires changing the behavior of staff and the \
climate, culture, and structures of the organization. As a result, \
implementing with fidelity and sustainability requires time, patience, \
discomfort, perseverance, and good leadership.')
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    "\tThe staff at your " + choice + " completed surveys to provide us with \
quantitative information. Additionally, optional focus groups with your \
" + choice + "'s \
staff may have given us additional qualitative information that we have used \
toward this final report. These focus groups facilitate our understanding of \
current practices and challenges that your staff face. Every organization has \
challenges; understanding exactly what yours are will help us know how and \
when to implement CPS to maximize its success."
)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tAfter you review this report, the Think:Kids team will discuss \
recommendations with ' + choice + ' leadership.  In some cases, immediate \
CPS implementation is recommended, and if so, a proposed comprehensive \
implementation and evaluation plan, timeline, and associated costs will be \
discussed.  In other cases, the Think:Kids team may recommend to leadership \
that particular readiness areas be addressed first, in order to maximize the \
benefits of implementation.'
)
document.add_page_break()

table = document.add_table(rows=1, cols=2, style='Normal Table')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Key:').bold = True
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Freq:').underline = True
row_cells[1].text = 'How many times a respondent responded in that way'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Percent:').underline = True
row_cells[1].text = 'The percent of the entire sample who responded in that \
way'
row_cells = table.add_row().cells
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Obs:').underline = True
row_cells[1].text = 'Number of respondents who completed this item or for \
whom the subscale could be calculated in a valid way'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Mean:').underline = True
row_cells[1].text = 'Average of all relevant items'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Std Dev:').underline = True
row_cells[1].text = 'Standard Deviation: Approximately 68% of responses fall \
within the mean plus or minus the standard deviation. For example, if the \
mean is 3 and standard deviation is 1, 68% of responses fall between 3+/-1, \
or between 2 and 4.'
row_cells = table.add_row().cells
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Min:').underline = True
row_cells[1].text = 'Minimum score for that item or category'
row_cells = table.add_row().cells
row_cells[0].paragraphs[0].add_run('Max:').underline = True
row_cells[1].text = 'Maximum score for that item or category'

document.add_picture('boxplotdescription.png', width=Inches(6.5))


def set_col_widths(table, first_col, other_cols):
    widths = (Inches(first_col), Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


set_col_widths(table, 1, 6)


# II. Quantitative Assessment
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tII.\t'
)
paragraph.add_run(
    'Quantitative Assessment'
).underline = True
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    'The CPS Adherence and Impact Measure (CPS-AIM)'
).bold = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tThe CPS-AIM inquires about factors related to CPS. This pre-training \
measurement was taken so that we can monitor how staff’s adherence to the CPS \
philosophy increases over time. We also hope that over time we will see \
your staff reporting reduced burnout.')
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Survey Responders:'
).bold = True
paragraph.add_run(
    '\n\t' + str(survey_count_aim) + ' total staff members responded to our \
CPS Adherence and Impact Measure Survey in a valid and reliable way. The \
quantitative analyses on the next few pages of this report are based on data \
collected from those respondents. The respondents are broken down by \
program, job role, and CPS training status as follows:\n'
)


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Staff Role Educators Table
if role_survey_count_aime > 0:
    make_table(
        'Staff Roles Educators',
        role_count_aime,
        role_dictionary_aime,
        role_survey_count_aime
    )


# Staff Role Clinical Table
if role_survey_count_aimc > 0:
    make_table(
        'Staff Roles Clinical',
        role_count_aimc,
        role_dictionary_aimc,
        role_survey_count_aimc
    )


# Staff Program Table
if sub_org_dict:
    document.add_paragraph('', style='Normal')
    make_table(
        'Staff Program',
        sub_org_count,
        sub_org_dict,
        sub_org_survey_count
    )


# Staff Training Table
document.add_paragraph('', style='Normal')
table = document.add_table(rows=1, cols=3, style='Normal Table')
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('CPS Training').bold = True
headers = ['Freq', 'Percent']
for count, header in enumerate(headers):
    hdr_cells[count + 1].paragraphs[0].add_run(header).underline = True
total = 0
for key, value in training_dictionary.items():
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = str(value)
    total += int(value)
    row_cells[2].text = str(
        round(float((value / training_total) * 100), 0)) + '%'
set_col_widths(table, 4, 1)


document.add_page_break()


def box_plot(title,
             df,
             sub_title='',
             dashed=False,
             dashed_place=4,
             max=7,
             min=1):
    """
    title = the large text at the top of the graph
    sub_title = the smaller text below the title
    df = the dataframe containing the data for the box plot
    """
    plt.figure()
    sns.set(style='whitegrid')
    sns.boxplot(data=df)
    plt.suptitle(title, fontsize=14)
    plt.title(sub_title, fontsize=10)
    plt.xticks(fontsize=12)
    plt.ylim(ymax=max, ymin=min)
    if dashed is True:
        plt.axhline(y=dashed_place, color='r', linestyle='--')
    plt.savefig('plt.png')
    document.add_picture('plt.png')


# The CPS Adherence and Impact Measure (CPS-AIM)
# Creating the CPS AIM Graph for Pre-Training
# Score the CPS AIM Educators
cps_aime_cols = []

# Getting only columns with CPS-AIM Educators Data
for i in cps_aim_df.columns:
    if 'cpsaime' in i:
        cps_aime_cols.append(i)

cps_aime_cols.append('record_id')

# Cleaning the CPS-AIM Educators columns to match the scoring program columns
new_cps_aime_columns = []

for i in range(32):
    new_cps_aime_columns.append('tkcot_' + str(i + 1))

# Question has been dropped from CPS-AIM
new_cps_aime_columns.remove('tkcot_11')
new_cps_aime_columns.append('record_id')

# Create a new DataFrame with the CPS-AIM Educators Data
cps_aime_df = cps_aim_df[cps_aime_cols]

# Change the column names of the new DataFrame to match scoring program names
cps_aime_df.columns = new_cps_aime_columns

# Drop out data for participants who missed 5 or more questions
cps_aime_df = cps_aime_df.dropna(thresh=5)

# Empty Results
results = []
# Scoring the CPS-AIM Educators
results = score.cps_aim_educator(cps_aime_df, results, 'record_id')

# Dropping out the Record ID column
cps_aime_df = cps_aime_df.drop(['record_id'], axis=1)

cps_aime_df = cps_aime_df.reset_index()
# Dropping the Index column
cps_aime_df = cps_aime_df.drop(['index'], axis=1)

# Creating lists of column names for each sub-scale
adherence_columns = []
perception_columns = []
burnout_columns = []
reverse_columns = []

subscales = {
    'Adherence to CPS Philosophy': {
        'Title': 'CPS-AIM Educators: Philosophy',
        'Columns': adherence_columns,
        'Description': cps_aim_subscales_text.iloc[1]['Text']
    },
    'Burnout': {
        'Title': 'CPS-AIM Educators: Burnout',
        'Columns': burnout_columns,
        'Description': cps_aim_subscales_text.iloc[2]['Text']
    }
}

# And a final column list for the CPS-AIM Educators
new_cps_aime_columns = []

for i in cps_aime_df.columns:
    for key, value in score.cps_aim_edu_items_dict.items():
        if key == i:
            new_cps_aime_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Adherence to CPS Philosophy':
                adherence_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Perception of Positive Impact':
                perception_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Burnout':
                burnout_columns.append(value['Trunc Question'])
            if value['Reverse'] == 'True':
                reverse_columns.append(value['Trunc Question'])

# Final CPS-AIM Educators Results DataFrame
cps_aime_df.columns = new_cps_aime_columns

# Re-order the columns by subscale
for i in reverse_columns:
    cps_aime_df[i].replace(
        [1, 2, 3, 4, 5, 6, 7], [7, 6, 5, 4, 3, 2, 1],
        inplace=True)

if not results.empty:
    paragraph = document.add_paragraph('', style='Normal')
    paragraph.add_run(cps_aim_subscales_text.iloc[0]['Text'])
    sns_results = results[results.columns[[1, 3]]]
    sns_results.columns = ['Philosophy', 'Burnout']
    box_plot(title='CPS Adherence and Impact Measure for Educational Staff',
             df=sns_results,
             sub_title='Pre-Training Score',
             dashed=True)

    # Adding the Summary: section, to be filled out later by Alisha
    paragraph = document.add_paragraph('', style='Normal')
    paragraph.add_run('Summary:').bold = True
    paragraph.add_run(' As a whole...')
    paragraph = document.add_paragraph('', style='Normal')
    for i in subscales.values():
        heat_map(
            df=cps_aime_df[i['Columns']],
            chart_title=i['Title'],
            min=1,
            max=7,
            height=8.5,
            width=4)
        table = document.add_table(rows=1, cols=2, style='Normal Table')
        cell = table.rows[0].cells[0].paragraphs[0].add_run(
        ).add_picture('plt.png')
        cell = table.rows[0].cells[1].paragraphs[0].add_run(i['Description'])


# Score the CPS AIM Systems
cps_aims_cols = []

# Getting only columns with CPS-AIM Systems Data
for i in cps_aim_df.columns:
    if 'cpsaims' in i:
        cps_aims_cols.append(i)

cps_aims_cols.append('record_id')

# Cleaning the CPS-AIM Systems columns to match the scoring program columns
new_cps_aims_columns = []

for i in range(36):
    new_cps_aims_columns.append('tkcot_' + str(i + 1))

# Questions have been dropped from CPS-AIM
new_cps_aims_columns.remove('tkcot_30')
new_cps_aims_columns.remove('tkcot_29')
new_cps_aims_columns.append('record_id')

# Create a new DataFrame with the CPS-AIM Systems Data
cps_aims_df = cps_aim_df[cps_aims_cols]

# Change the column names of the new DataFrame to match scoring program names
cps_aims_df.columns = new_cps_aims_columns

# Drop out data for participants who missed 5 or more questions
cps_aims_df = cps_aims_df.dropna(thresh=5)

# Empty Results
results = []
# Scoring the CPS-AIM Educators
results = score.cps_aim_systems(cps_aims_df, results, 'record_id')

# Dropping out the Record ID column
cps_aims_df = cps_aims_df.drop(['record_id'], axis=1)

cps_aims_df = cps_aims_df.reset_index()
# Dropping the Index column
cps_aims_df = cps_aims_df.drop(['index'], axis=1)

# Creating lists of column names for each sub-scale
adherence_columns = []
perception_columns = []
burnout_columns = []
reverse_columns = []

subscales = {
    'Adherence to CPS Philosophy': {
        'Title': 'CPS-AIM Clinical: Philosophy',
        'Columns': adherence_columns,
        'Description': cps_aim_subscales_text.iloc[1]['Text']
    },
    'Burnout': {
        'Title': 'CPS-AIM Clinical: Burnout',
        'Columns': burnout_columns,
        'Description': cps_aim_subscales_text.iloc[2]['Text']
    }
}

# And a final column list for the CPS-AIM Systems
new_cps_aims_columns = []

for i in cps_aims_df.columns:
    for key, value in score.cps_aim_sys_items_dict.items():
        if key == i:
            new_cps_aims_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Adherence to CPS Philosophy':
                adherence_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Perception of Positive Impact':
                perception_columns.append(value['Trunc Question'])
            if value['Subscale'] == 'Burnout':
                burnout_columns.append(value['Trunc Question'])
            if value['Reverse'] == 'True':
                reverse_columns.append(value['Trunc Question'])


# Final CPS-AIM Educators Results DataFrame
cps_aims_df.columns = new_cps_aims_columns

# Re-order the columns by subscale
for i in reverse_columns:
    cps_aims_df[i].replace(
        [1, 2, 3, 4, 5, 6, 7], [7, 6, 5, 4, 3, 2, 1],
        inplace=True)

if not results.empty:
    paragraph = document.add_paragraph('', style='Normal')
    paragraph.add_run(cps_aim_subscales_text.iloc[0]['Text'])
    sns_results = results[results.columns[[1, 3]]]
    sns_results.columns = ['Philosophy', 'Burnout']
    box_plot(title='CPS Adherence and Impact Measure for Clinical Staff',
             df=sns_results,
             sub_title='Pre-Training Score',
             dashed=True)

    # Adding the Summary: section, to be filled out later by Alisha
    paragraph = document.add_paragraph('', style='Normal')
    paragraph.add_run('Summary:').bold = True
    paragraph.add_run(' As a whole...')
    paragraph = document.add_paragraph('', style='Normal')
    for i in subscales.values():
        heat_map(
            df=cps_aims_df[i['Columns']],
            chart_title=i['Title'],
            min=1,
            max=7,
            height=8.5,
            width=4)
        table = document.add_table(rows=1, cols=2, style='Normal Table')
        cell = table.rows[0].cells[0].paragraphs[0].add_run(
        ).add_picture('plt.png')
        cell = table.rows[0].cells[1].paragraphs[0].add_run(i['Description'])


# Pull REDcap Token depending on user choice
# Part 2!
# "CPS Readiness Assessment (1-Part)"
if choice == 'school':
    tk.readiness_for_schools()
elif choice == 'system':
    tk.readiness_for_systems()

project = Project(tk.api_url, tk.api_token)
data_readiness = project.export_records(format='df')
metadata_readiness = project.export_metadata(format='df')

organization_dictionary_readiness = meta_dict(
    metadata_readiness, 'organization')

# Download the corresponding Readiness Survey Data
print("The organizations listed in the CPS Readiness Assessment for \
" + choice.title() + "s (1-Part) are as follows.")
for key, value in organization_dictionary_readiness.items():
    print(str(key) + '\t' + value)
org_number_readiness = int(input("Please enter the Organizations \
Number to generate a report on.\n\t"))
print('Thanks!\n\n\n')

# Retrieving and sorting the Readiness Data from REDcap
data_readiness.reset_index(inplace=True)
org_data_readiness = data_readiness[
    data_readiness['organization'] == org_number_readiness
]
org_data_readiness = org_data_readiness.dropna(thresh=15)


# Sub School
sub_org_variable_name = []
for x, i in enumerate(metadata_readiness['branching_logic']):
    if '[' + cps_readiness_org + '] = ' + "'" + str(
            org_number_readiness) + "'" in str(i):
        sub_org_variable_name.append(metadata_readiness.index[x])
sub_org_dict = {}
for i in sub_org_variable_name:
    sub_org_variable_name = i
    sub_org_dict = meta_dict(metadata_readiness, i)
    sub_org_count = org_data_readiness[i].value_counts().to_dict()
    sub_org_survey_count = org_data_readiness[i].count()

if breakdown == 2:
    for key, value in sub_org_dict.items():
        print(str(key) + '\t' + value)
    org_number_readiness = int(input("Please enter the program \
    number to generate a report on.\n\t"))
    print('Thanks!\n\n\n')

# Retrieving and sorting the Readiness Data from REDcap
    org_data_readiness = org_data_readiness[
        org_data_readiness[sub_org_variable_name] == org_number_readiness]

# Remove participants who did not answer at least 15 questions

# Replace any N/A answers (9) with np.nan so that the data is not skewed
sub_org_column = org_data_readiness[sub_org_variable_name]
org_data_readiness = org_data_readiness.replace(9, np.nan)
org_data_readiness[sub_org_variable_name] = sub_org_column


# Reverse the Reverse scored columns
score.reversed_readiness_df(org_data_readiness)

# Variables from Readiness Data Set

# Get a count of how many participants responded to the survey and answered
# at least 15 questions.
survey_count_readiness = org_data_readiness.record_id.count()

# Role variables
role_dictionary_readiness = meta_dict(metadata_readiness, 'rolestaff')
role_count_readiness = (
    org_data_readiness['rolestaff'].value_counts().to_dict())
role_survey_count_readiness = org_data_readiness['rolestaff'].count()

# Sub Organization variables
sub_org_dictionary_readiness = (
    meta_dict(metadata_readiness, sub_org_variable_name)
)

sub_org_count_readiness = (
    org_data_readiness[sub_org_variable_name].value_counts().to_dict()
)
sub_org_survey_count_readiness = (
    org_data_readiness[sub_org_variable_name].count()
)

# Years at Organization Dictionary - Readiness
years_at_org = meta_dict(metadata_readiness, 'yearsatorg')
years_at_org_count = (
    org_data_readiness['yearsatorg'].value_counts().to_dict())
years_at_org_total_count = org_data_readiness['yearsatorg'].count()

# Training Dictionary
training_dictionary = meta_dict(metadata_readiness, 'training')
training_count = (org_data_readiness['training'].value_counts().to_dict())
training_total_count = org_data_readiness['training'].count()

# CPS Readiness Assessment
add_document_text_from_excel(readiness_survey_text)
document.add_page_break()

# Staff Role Table Readiness
make_table(
    'Staff Role',
    role_count_readiness,
    role_dictionary_readiness,
    role_survey_count_readiness
)

# Staff Program Table Readiness

if sub_org_dict:
    document.add_paragraph('', style='Normal')
    make_table(
        'Staff Program',
        sub_org_count,
        sub_org_dict,
        sub_org_survey_count
    )

# Years at Organization Table Readiness
document.add_paragraph('', style='Normal')

make_table(
    'Years at Organization',
    years_at_org_count,
    years_at_org,
    years_at_org_total_count
)

document.add_page_break()

# Readiness Survey Results, for All Employees
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run('Readiness Survey Results, for All Employees').bold = True

# Score the Readiness Survey
# Score the Educators Version
results_educators = []
results_educators = score.cps_readiness(
    org_data_readiness,
    results_educators,
    'record_id',
    'staff')
results_educators = results_educators.dropna(thresh=3)

# Score the Admin Version
results_admin = []
results_admin = score.cps_readiness(
    org_data_readiness,
    results_admin,
    'record_id',
    'admin')
results_admin = results_admin.dropna(thresh=3)

frames = [results_educators, results_admin]
results = pd.concat(frames)

# Add the scored columns back into the main DataFrame.
org_data_readiness = pd.merge(
    left=org_data_readiness,
    right=results,
    left_on='record_id',
    right_on='id')
# CPS Readiness Measure Score Variables.
fit_dict_admin = {
    'admin1': 'Policies need improvement',
    'admin2': 'CPS is improvement',
    'admin5': 'CPS too hard (reversed)',
    'admin6': 'Leaders want all in',
    'admin7': 'CPS consistent with values',
    'admin8': 'CPS consistent with practice'
}
fit_dict_staff = {
    'staff1': 'Policies need improvement',
    'staff2': 'CPS is improvement',
    'staff5': 'CPS too hard (reversed)',
    'staff6': 'Leaders want all in',
    'staff7': 'CPS consistent with values',
    'staff8': 'CPS consistent with practice'
}
capacity_dict_admin = {
    'admin3': 'Leaders are committed',
    'admin4': 'Internal CPS team',
    'admin9': 'We encourage innovation',
    'admin10': 'Staff want to learn more',
    'admin11': 'Staff supported by leaders',
    'admin12': 'Staff communicate well',
    'admin13': 'Too many interventions (reversed)',
    'admin14': 'Financially committed',
    'admin15': 'Time to brainstorm',
    'admin16': 'Time to pull youth aside'
}
capacity_dict_staff = {
    'staff3': 'Leaders are committed',
    'staff4': 'Internal CPS team',
    'staff9': 'We encourage innovation',
    'staff10': 'Staff want to learn more',
    'staff11': 'Staff supported by leaders',
    'staff12': 'Staff communicate well',
    'staff13': 'Too many interventions (reversed)',
    'staff14': 'Receive supervision or mentorship',
    'staff15': 'Time to brainstorm',
    'staff16': 'Time to pull youth aside'
}

# BIG DICTIONARY OF READINESS INFORMATION!!!
readiness_data = {
    'Fit': {
        'All Staff': {
            'Overall Stats': stats_dict(
                org_data_readiness, 'readiness_fit_mean'),
        },
        'Administration': {
            'Overall': stats_dict(
                results_admin, 'readiness_fit_mean')
        },
        'Staff': {
            'Overall': stats_dict(
                results_educators, 'readiness_fit_mean')
        }
    },
    'Capacity': {
        'All Staff': {
            'Overall Stats': stats_dict(
                org_data_readiness, 'readiness_capacity_mean'),
        },
        'Administration': {
            'Overall': stats_dict(
                results_admin, 'readiness_capacity_mean')
        },
        'Staff': {
            'Overall': stats_dict(
                results_educators, 'readiness_capacity_mean')
        }
    }
}


for key, value in fit_dict_admin.items():
    readiness_data['Fit']['Administration'][key] = {
        'Full Question': cell_content(metadata_readiness, key, 'field_label'),
        'Trunc Question': value,
        'Stats': stats_dict(org_data_readiness, key)
    }

for key, value in fit_dict_staff.items():
    readiness_data['Fit']['Staff'][key] = {
        'Full Question': cell_content(metadata_readiness, key, 'field_label'),
        'Trunc Question': value,
        'Stats': stats_dict(org_data_readiness, key)
    }

for key, value in capacity_dict_admin.items():
    readiness_data['Capacity']['Administration'][key] = {
        'Full Question': cell_content(metadata_readiness, key, 'field_label'),
        'Trunc Question': value,
        'Stats': stats_dict(org_data_readiness, key)
    }

for key, value in capacity_dict_staff.items():
    readiness_data['Capacity']['Staff'][key] = {
        'Full Question': cell_content(metadata_readiness, key, 'field_label'),
        'Trunc Question': value,
        'Stats': stats_dict(org_data_readiness, key)
    }

# Admin
admin_readiness_results = {
    'Fit': stats_dict(
        results_admin, 'readiness_fit_mean'),
    'Capacity': stats_dict(
        results_admin, 'readiness_capacity_mean')
}

# Educators
readiness_results = {
    'Fit': stats_dict(
        results_educators, 'readiness_fit_mean'),
    'Capacity': stats_dict(
        results_educators, 'readiness_capacity_mean')
}

# All
all_readiness_results = {
    'Fit': stats_dict(
        org_data_readiness, 'readiness_fit_mean'),
    'Capacity': stats_dict(
        org_data_readiness, 'readiness_capacity_mean')
}


def set_col_widths(table, first_col, other_cols):
    widths = (
        Inches(first_col),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols),
        Inches(other_cols))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Tables describing statistics for all staff
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tPossible responses for readiness items range from 1 (Strongly Disagree)\
 to 5 (Strongly Agree). In the graph below, the dashed horizontal line \
indicates a score of 3, which is “Not Sure.”')
sns_results = org_data_readiness[org_data_readiness.columns[-2:]]
sns_results.columns = ['Fit', 'Capacity']
box_plot(title='Readiness Measure',
         df=sns_results,
         dashed=True,
         dashed_place=3,
         max=5)
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Summary:  Overall, staff at ' + org + ' are…\
\n\nThis spread can be seen in more detail below.')

document.add_page_break()

# *** Alisha might want to combine these two tables? ***
# Tables describing statistics for Staff
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run('Readiness Survey Results Staff').bold = True
readiness_tables(readiness_results)

# Tables describing statistics for Administration
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run(
    'Readiness Survey Results Leadership/Administration'
).bold = True
readiness_tables(admin_readiness_results)

paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Summary: Across roles, staff at are…')

document.add_page_break()

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Detail View: Fit and Capacity'
).bold = True

count_threshold = int(input('Under what number of staff in a particular \
role or program will we exclude staff from the graphs?\n\t'))
print('Thanks!\n\n\n')

# Staff Roles

roles_to_be_cut = []

for key, value in role_count_readiness.items():
    if value < count_threshold:
        roles_to_be_cut.append(key)

# Problem with picture which category is under threshold... Think On THIS!!
role_count_df = org_data_readiness
count_role_graphs = []
for index in role_count_df.index:
    if role_count_df.loc[index, 'rolestaff'] in roles_to_be_cut:
        role_count_df.drop(index, inplace=True)
    else:
        count_role_graphs.append(role_count_df.loc[index, 'rolestaff'])

count_role_graphs = int(np.ceil(len(set(count_role_graphs)) / 2))
count_role_graphs = int(count_role_graphs * 2)

unique_rolestaff = role_count_df['rolestaff'].nunique()
unique_rolestaff = int(np.ceil(unique_rolestaff / 2))
role_count_df['rolestaff'] = role_count_df['rolestaff'].map(
    role_dictionary_readiness)

paragraph = document.add_paragraph('\t', style='Normal')
paragraph.add_run(
    'In order to preserve confidentiality, categories where there were less \
than ' + str(count_threshold) + ' participants who answered the survey \
were excluded from the following graphs.')

histogram_by(
    title='Fit by Role',
    df=role_count_df,
    df_column='readiness_fit_mean',
    sort_by='rolestaff',
    height=count_role_graphs,
    columns=unique_rolestaff)
histogram_by(
    title='Capacity by Role',
    df=role_count_df,
    df_column='readiness_capacity_mean',
    sort_by='rolestaff',
    height=count_role_graphs,
    columns=unique_rolestaff)

# Programs
histograms_by_program = input('Would we like a breakdown of the histograms by \
program in addition to the breakdown by role?\n1\tYes\n2\tNo\n\t')
print('Thanks!')


if histograms_by_program == '1':
    program_count_df = org_data_readiness
    count_programs_graphs = []
    programs_to_be_cut = []
    for key, value in sub_org_count.items():
        if value < count_threshold:
            programs_to_be_cut.append(key)

    for index in program_count_df.index:
        if program_count_df.loc[
            index, sub_org_variable_name
        ] in programs_to_be_cut:
            program_count_df.drop(index, inplace=True)
        else:
            count_programs_graphs.append(
                program_count_df.loc[index, sub_org_variable_name])

    count_programs_graphs = int(np.ceil(len(set(count_programs_graphs)) / 2))
    count_programs_graphs = int(count_programs_graphs * 2)

    unique_programs = program_count_df[sub_org_variable_name].nunique()
    unique_programs = int(np.ceil(unique_programs / 2))
    program_count_df[
        sub_org_variable_name
    ] = program_count_df[sub_org_variable_name].map(
        sub_org_dictionary_readiness)

    histogram_by(
        title='Fit by Program',
        df=program_count_df,
        df_column='readiness_fit_mean',
        sort_by=sub_org_variable_name,
        height=count_programs_graphs,
        columns=unique_programs)
    histogram_by(
        title='Capacity by Program',
        df=program_count_df,
        df_column='readiness_capacity_mean',
        sort_by=sub_org_variable_name,
        height=count_programs_graphs,
        columns=unique_programs)

# ADD STAFF AND ADMIN TABLES
paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Detailed Report of ').bold = True
p = paragraph.add_run('Staffs’')
p.bold = True
p.underline = True
paragraph.add_run(' Responses by Item:').bold = True

detailed_report('staff')

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\t* Items have been truncated here to conserve space. See Appendix for \
item wording.'
).bold = True

paragraph = document.add_paragraph('\n', style='Normal')
paragraph.add_run('Detailed Report of ').bold = True
p = paragraph.add_run('Leadership/Administrators’')
p.bold = True
p.underline = True
paragraph.add_run(' Responses by Item:').bold = True

detailed_report('admin')

paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\t* Items have been truncated here to conserve space. See Appendix for \
item wording.'
).bold = True
document.add_page_break()

# Heat Maps
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run('\tHeat Maps\t').bold = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'The following is a graphical representation of each readiness survey \
item (columns) rated by each respondent (rows). Items have been truncated to \
save space; see the appendix for original item wording. Scores have been \
reversed when necessary so that darker colors always indicate better \
readiness. Thus, columns with a lot of beige or light green indicate \
readiness areas in need of improvement.'
)
paragraph = document.add_paragraph('', style='Body Text')


readiness_columns_admin = []
readiness_columns_admin_trunc = []
readiness_columns_staff = []
readiness_columns_staff_trunc = []

for key, value in readiness_data.items():
    for key, value in value.items():
        if key == 'Staff' or key == 'Administration':
            for key, value in value.items():
                if 'admin' in key:
                    readiness_columns_admin.append(key)
                    for key, value in value.items():
                        if 'Trunc Question' in key:
                            readiness_columns_admin_trunc.append(value)
                elif 'staff' in key:
                    readiness_columns_staff.append(key)
                    for key, value in value.items():
                        if 'Trunc Question' in key:
                            readiness_columns_staff_trunc.append(value)


admin = org_data_readiness[readiness_columns_admin].dropna(thresh=12)
admin_columns = []
for i in readiness_columns_admin_trunc:
    new = i.split(' ')
    start = []
    end = []
    count = 0
    for i in new:
        count += 1
        if count < 3:
            start.append(i)
        if count == 2:
            start = ' '.join(start)
        if count > 2:
            end.append(i)
    end = ' '.join(end)
    new_value = start + '\n' + end
    admin_columns.append(new_value)
admin = admin.reset_index().drop(['index'], axis=1)
admin.columns = admin_columns


staff = org_data_readiness[readiness_columns_staff].dropna(thresh=10)
staff_columns = []
for i in readiness_columns_staff_trunc:
    new = i.split(' ')
    start = []
    end = []
    count = 0
    for i in new:
        count += 1
        if count < 3:
            start.append(i)
        if count == 2:
            start = ' '.join(start)
        if count > 2:
            end.append(i)
    end = ' '.join(end)
    new_value = start + '\n' + end
    staff_columns.append(new_value)
staff = staff.reset_index().drop(['index'], axis=1)
staff.columns = staff_columns

if not staff.empty:
    heat_map(
        df=staff,
        chart_title='Readiness Survey Responses Heat Map for Staff',
        min=1,
        max=5,
        height=7,
        width=6.5
    )
    paragraph.add_run().add_picture('plt.png')
    document.add_page_break()

if not admin.empty:
    heat_map(
        df=admin,
        chart_title='Readiness Survey Responses Heat Map for Admin',
        min=1,
        max=5,
        height=6.5,
        width=6.5
    )
    paragraph = document.add_paragraph('', style='Body Text')
    paragraph.add_run().add_picture('plt.png')
    document.add_page_break()

# III. Qualitative Assessment
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run(
    '\tIII.\t'
)
paragraph.add_run(
    'Qualitative Assessment'
).underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Summary of Need:'
).bold = True
paragraph.add_run('\n')
paragraph.add_run(
    '\nReadiness strengths:'
).bold = True
paragraph.add_run('\n')
paragraph.add_run(
    '\nReadiness areas in need of improvement:'
).bold = True
paragraph.add_run('\n')
paragraph = document.add_paragraph(
    '\nReadiness Area: Fit', style='Normal'
)
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph(
    '\nReadiness Area: Capacity', style='Normal'
)
document.add_paragraph(
    '', style='List Bullet 2'
)
document.add_page_break()


# IV. CPS Readiness Summary

paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run('\tIV.\t')
paragraph.add_run('CPS Readiness Summary').underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Summary of Need:'
).bold = True
paragraph.add_run('\n')
paragraph.add_run(
    '\n' + org + ' is...'
)
paragraph.add_run(
    '\nOn the whole, ' + org + ' staff...'
)
paragraph.add_run('\n')
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\nReadiness Strengths and Areas for Improvement'
).bold = True
paragraph.add_run(
    '\nBased on the quantitative results of the readiness assessment, \
' + org + ' appears to be <options: very well positioned, well \
positioned, not yet ready> to implement CPS.')
paragraph.add_run(
    '\nReadiness strengths:'
).italic = True
document.add_paragraph(
    '', style='List Bullet 2'
)
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\nReadiness area for improvement:'
).italc = True
document.add_paragraph(
    '', style='List Bullet 2'
)
document.add_page_break()

# Creating the Readiness Measure Template

table = document.add_table(rows=1, cols=3, style='Normal Table')
hdr_cells = table.rows[0].cells
hdr_cells[1].paragraphs[0].add_run(
    'Score:    0=Not at all    1=Partially    2=Definitely'
)

for row in total_readiness_table_text.index:
    row_cells = table.add_row().cells
    item = str(total_readiness_table_text.iloc[row]['Item'])
    body = str(total_readiness_table_text.iloc[row]['Body'])
    score = str(total_readiness_table_text.iloc[row]['Score'])
    style = str(total_readiness_table_text.iloc[row]['Style'])
    if style == 'bold':
        row_cells[0].paragraphs[0].add_run(item).bold = True
        row_cells[1].paragraphs[0].add_run(body).bold = True
        row_cells[2].paragraphs[0].add_run(score).bold = True
    else:
        row_cells[0].paragraphs[0].add_run(item)
        row_cells[1].paragraphs[0].add_run(body)
        row_cells[2].paragraphs[0].add_run(score)


for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(10)


def set_col_widths(table, first_col, middle_col, last_col):
    widths = (
        Inches(first_col),
        Inches(middle_col),
        Inches(last_col))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


set_col_widths(table, 1.25, 8, .5)

document.add_page_break()

# V. Recommendations
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run('\tV.\t')
paragraph.add_run('Recommendations').underline = True
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    '\tBased upon the complete results of this readiness assessment, our \
Readiness Team feels that ' + org + ' is in excellent shape to continue to \
the next phase of implementation. Our recommendations are as follows:'
)
paragraph.add_run('\n')

paragraph = document.add_paragraph(
    '\nPrior to Training', style='Normal'
).bold = True
document.add_paragraph('', style='List Bullet 2')
paragraph = document.add_paragraph(
    '\nTraining and Coaching', style='Normal'
).bold = True
document.add_paragraph('', style='List Bullet 2')
paragraph = document.add_paragraph(
    '\nMoving Toward Sustainability', style='Normal'
).bold = True
document.add_paragraph('', style='List Bullet 2')
document.add_page_break()

# Appendix
paragraph = document.add_paragraph('', style='Body Text')
paragraph.add_run('Appendix').underline = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run(
    'Organization-Wide Implementation Readiness Survey'
).bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run('\nFOR STAFF:')

question_columns_staff = []

questions_staff = []
for i in metadata_readiness.index:
    if 'staff' in i:
        question_columns_staff.append(i)
        question = cell_content(metadata_readiness, i, 'field_label')
        for i in question:
            questions_staff.append(i)
questions_staff = questions_staff[1:]
question_columns_staff = question_columns_staff[1:]


add_survey(questions_staff)
document.add_page_break()
paragraph = document.add_paragraph('', style='Normal')
paragraph.add_run('\nFOR LEADERS/ADMINISTRATION:')

question_columns_admin = []
questions_admin = []
for i in metadata_readiness.index:
    if 'admin' in i:
        question_columns_admin.append(i)
        question = cell_content(metadata_readiness, i, 'field_label')
        for i in question:
            questions_admin.append(i)
add_survey(questions_admin)

os.remove('plt.png')

document.save(org + ' Readiness Report.docx')
print('The Readiness Report has been generated and saved!')
